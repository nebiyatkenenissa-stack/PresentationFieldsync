from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(30, 58, 95)
    hs.font.name = 'Calibri'

NAVY = '1E3A5F'
GREEN = '2E7D32'
RED = 'C62828'
ORANGE = 'E65100'
BLUE = '1565C0'

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells_data, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        if header:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_shading(cell, NAVY)
    return row

def make_table(doc, headers, rows):
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, headers, header=True)
    for row_data in rows:
        add_table_row(table, row_data)
    doc.add_paragraph()
    return table

def add_method_badge(doc, method, path, desc):
    p = doc.add_paragraph()
    method_colors = {'GET': GREEN, 'POST': BLUE, 'PUT': ORANGE, 'DELETE': RED}
    color = method_colors.get(method, NAVY)
    run = p.add_run(f'  {method}  ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(255, 255, 255)
    r, g, b = int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
    run.font.highlight_color = None
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    run._r.get_or_add_rPr().append(shading)
    run2 = p.add_run(f'  {path}')
    run2.bold = True
    run2.font.size = Pt(11)
    run2.font.name = 'Consolas'
    run3 = p.add_run(f'    {desc}')
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(100, 100, 100)

def add_endpoint_section(doc, method, path, desc, params=None, req_body=None, req_example=None, res_example=None, errors=None):
    add_method_badge(doc, method, path, desc)

    if params:
        doc.add_paragraph('Request Parameters:', style='List Bullet')
        make_table(doc, ['Name', 'Type', 'Required', 'Description'], params)

    if req_body:
        p = doc.add_paragraph()
        run = p.add_run('Request Body:')
        run.bold = True
        run.font.size = Pt(10)
        make_table(doc, ['Field', 'Type', 'Required', 'Description'], req_body)

    if req_example:
        p = doc.add_paragraph()
        run = p.add_run('Example Request:')
        run.bold = True
        run.font.size = Pt(10)
        p = doc.add_paragraph()
        run = p.add_run(req_example)
        run.font.name = 'Consolas'
        run.font.size = Pt(8)

    if res_example:
        p = doc.add_paragraph()
        run = p.add_run('Example Response:')
        run.bold = True
        run.font.size = Pt(10)
        p = doc.add_paragraph()
        run = p.add_run(res_example)
        run.font.name = 'Consolas'
        run.font.size = Pt(8)

    if errors:
        p = doc.add_paragraph()
        run = p.add_run('Error Responses:')
        run.bold = True
        run.font.size = Pt(10)
        make_table(doc, ['Code', 'Condition'], errors)

    doc.add_paragraph('---')

# =============================================
# COVER PAGE
# =============================================
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FieldSync')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(30, 58, 95)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('API Documentation')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(79, 195, 247)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('RESTful API Reference for FieldSync Backend Services')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Version 1.0.0\n{datetime.date.today().strftime("%B %Y")}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 116, 139)

doc.add_page_break()

# =============================================
# INTRODUCTION
# =============================================
doc.add_heading('Introduction', level=1)

doc.add_heading('Purpose', level=2)
doc.add_paragraph(
    'This document provides a comprehensive overview of the FieldSync REST API, including '
    'all available endpoints, request/response formats, authentication requirements, error '
    'handling, and technical specifications intended to assist developers and system integrators '
    'in utilizing the API effectively.'
)

doc.add_heading('General Information', level=2)
make_table(doc, ['Field', 'Value'], [
    ['API Base URL', 'http://localhost:5001 (development) / https://api.fieldsync.gov (production)'],
    ['API Version', 'v1 (1.0.0)'],
    ['Contact Email', 'api-support@fieldsync.gov'],
    ['Authentication Method', 'JWT Bearer Token (JSON Web Token)'],
    ['Content Type', 'application/json (all requests and responses)'],
    ['Total Endpoints', '47 endpoints across 13 resource modules'],
    ['Architecture', 'RESTful, MVC pattern (Express 5 + TypeScript)'],
])

doc.add_paragraph()
doc.add_heading('Authentication', level=2)
doc.add_paragraph(
    'The FieldSync API uses JWT (JSON Web Token) based authentication. All API endpoints '
    'require a valid JWT token in the Authorization header, except for the login endpoint. '
    'The token is obtained by calling the /api/auth/login endpoint with valid email and password '
    'credentials. The server validates the token on each request and extracts the user identity '
    'for authorization purposes.'
)
doc.add_paragraph(
    'Password Security: All passwords are hashed using bcrypt with 10 salt rounds before storage. '
    'Plaintext passwords are never transmitted to or stored by the server after initial authentication.'
)

doc.add_heading('Authentication Example', level=2)
p = doc.add_paragraph()
run = p.add_run(
    '# Step 1: Login to obtain token\n'
    'POST /api/auth/login\n'
    'Content-Type: application/json\n\n'
    '{\n'
    '  "email": "officer@fieldsync.gov",\n'
    '  "password": "SecurePass123"\n'
    '}\n\n'
    '# Response:\n'
    '{\n'
    '  "success": true,\n'
    '  "user": {\n'
    '    "id": "usr_001",\n'
    '    "employee_id": "FO-001",\n'
    '    "name": "Abebe Kebede",\n'
    '    "email": "officer@fieldsync.gov",\n'
    '    "role": "field_officer",\n'
    '    "status": "active"\n'
    '  }\n'
    '}\n\n'
    '# Step 2: Use token in subsequent requests\n'
    'GET /api/citizens\n'
    'Authorization: Bearer eyJhbGciOiJIUzI1NiIs...\n'
    'Content-Type: application/json'
)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_paragraph()
doc.add_heading('Error Codes', level=2)
make_table(doc, ['Code', 'Message', 'Description'], [
    ['200', 'OK', 'Request succeeded.'],
    ['201', 'Created', 'Resource successfully created.'],
    ['400', 'Bad Request', 'Missing required fields or invalid request body format.'],
    ['401', 'Unauthorized', 'Invalid credentials, missing token, or inactive account.'],
    ['404', 'Not Found', 'Requested resource does not exist.'],
    ['409', 'Conflict', 'Duplicate resource (e.g., duplicate email or citizen name).'],
    ['500', 'Internal Server Error', 'Unexpected server-side error. Check server logs for details.'],
])

doc.add_page_break()

# =============================================
# RESOURCES / ENDPOINTS
# =============================================
doc.add_heading('Resources / Endpoints', level=1)

# ---- AUTH ----
doc.add_heading('Authentication (/api/auth)', level=2)

add_endpoint_section(doc, 'POST', '/api/auth/login', 'Authenticate a user and receive a JWT token.',
    req_body=[
        ['email', 'string', 'Yes', 'User email address'],
        ['password', 'string', 'Yes', 'User password'],
    ],
    req_example='POST /api/auth/login\nContent-Type: application/json\n\n{"email": "user@fieldsync.gov", "password": "SecurePass123"}',
    res_example='{"success": true, "user": {"id": "usr_001", "employee_id": "FO-001", "name": "Abebe Kebede", "email": "user@fieldsync.gov", "role": "field_officer", "status": "active"}}',
    errors=[['400', 'Missing email or password'], ['401', 'Invalid credentials or inactive account'], ['500', 'Server error']]
)

add_endpoint_section(doc, 'POST', '/api/auth/change-password', 'Change the authenticated user\'s password.',
    req_body=[
        ['email', 'string', 'Yes', 'User email address'],
        ['currentPassword', 'string', 'Yes', 'Current password for verification'],
        ['newPassword', 'string', 'Yes', 'New password to set'],
    ],
    req_example='POST /api/auth/change-password\nContent-Type: application/json\n\n{"email": "user@fieldsync.gov", "currentPassword": "OldPass", "newPassword": "NewPass123"}',
    res_example='{"success": true, "message": "Password changed successfully"}',
    errors=[['400', 'Missing required fields'], ['401', 'Current password incorrect'], ['404', 'User not found'], ['500', 'Server error']]
)

# ---- USERS ----
doc.add_heading('Users (/api/users)', level=2)

add_endpoint_section(doc, 'GET', '/api/users', 'Retrieve a list of all users. Managers see all users; Supervisors see their team.',
    res_example='[{"id": "usr_001", "employee_id": "FO-001", "name": "Abebe Kebede", "email": "a@fieldsync.gov", "role": "field_officer", "region": "Addis Ababa", "status": "active"}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'GET', '/api/users/:id', 'Retrieve a single user by ID.',
    params=[['id', 'string', 'Yes', 'User ID (path parameter)']],
    res_example='{"id": "usr_001", "employee_id": "FO-001", "name": "Abebe Kebede", "email": "a@fieldsync.gov", "role": "field_officer", "status": "active", "phone": "+251911000000"}',
    errors=[['404', 'User not found'], ['500', 'Server error']]
)

add_endpoint_section(doc, 'GET', '/api/users/supervisors-by-woreda/:woredaId', 'Retrieve supervisors by woreda (district) ID.',
    params=[['woredaId', 'integer', 'Yes', 'Woreda location ID (path parameter)']],
    res_example='[{"id": "usr_010", "employee_id": "SUP-001", "name": "Fatuma Ali", "email": "f@fieldsync.gov"}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'POST', '/api/users', 'Create a new user account.',
    req_body=[
        ['id', 'string', 'Yes', 'Unique user ID'],
        ['employeeId', 'string', 'Yes', 'Employee identifier'],
        ['name', 'string', 'Yes', 'Full name'],
        ['email', 'string', 'Yes', 'Unique email address'],
        ['password', 'string', 'No', 'Initial password (auto-generated if omitted)'],
        ['role', 'string', 'Yes', 'Role: manager, supervisor, or field_officer'],
        ['region', 'string', 'No', 'Assigned region'],
        ['supervisorId', 'string', 'No', 'Assigned supervisor ID'],
        ['status', 'string', 'No', 'Account status (default: active)'],
        ['phone', 'string', 'No', 'Phone number'],
        ['country_id', 'integer', 'No', 'Country location ID'],
        ['region_id', 'integer', 'No', 'Region location ID'],
        ['zone_id', 'integer', 'No', 'Zone location ID'],
        ['woreda_id', 'integer', 'No', 'Woreda location ID'],
        ['kebele_id', 'integer', 'No', 'Kebele location ID'],
    ],
    res_example='{"id": "usr_002", "name": "New User", "email": "new@fieldsync.gov", "role": "field_officer", "temporaryPassword": "Temp8392"}',
    errors=[['400', 'Missing email'], ['409', 'Duplicate employee_id'], ['500', 'Server error']]
)

add_endpoint_section(doc, 'POST', '/api/users/resend-credentials', 'Resend login credentials (email and temporary password) to a user.',
    req_body=[['email', 'string', 'Yes', 'User email address']],
    res_example='{"ok": true, "temporaryPassword": "Temp8392", "email": "user@fieldsync.gov", "note": "Credentials resent"}',
    errors=[['404', 'User not found'], ['500', 'Server error']]
)

add_endpoint_section(doc, 'POST', '/api/users/:id/photo', 'Upload a profile photo for a user.',
    params=[['id', 'string', 'Yes', 'User ID (path parameter)']],
    req_example='POST /api/users/usr_001/photo\nContent-Type: multipart/form-data\n\nprofilePhoto: <binary file data>',
    res_example='{"profilePhoto": "/uploads/profile/usr_001.jpg"}',
    errors=[['404', 'User not found'], ['500', 'Upload failed']]
)

add_endpoint_section(doc, 'PUT', '/api/users/:id', 'Update user profile information.',
    params=[['id', 'string', 'Yes', 'User ID (path parameter)']],
    req_body=[
        ['name', 'string', 'No', 'Updated name'],
        ['email', 'string', 'No', 'Updated email'],
        ['phone', 'string', 'No', 'Updated phone number'],
        ['role', 'string', 'No', 'Updated role'],
        ['status', 'string', 'No', 'Updated status (active/inactive)'],
        ['region', 'string', 'No', 'Updated region'],
        ['supervisorId', 'string', 'No', 'Updated supervisor'],
    ],
    res_example='{"id": "usr_001", "name": "Updated Name", "email": "a@fieldsync.gov", "role": "field_officer", "status": "active"}',
    errors=[['404', 'User not found'], ['409', 'Duplicate employee_id'], ['500', 'Server error']]
)

add_endpoint_section(doc, 'DELETE', '/api/users/:id', 'Deactivate/delete a user account.',
    params=[['id', 'string', 'Yes', 'User ID (path parameter)']],
    res_example='{"message": "User deleted successfully"}',
    errors=[['404', 'User not found'], ['500', 'Server error']]
)

# ---- CITIZENS ----
doc.add_heading('Citizens (/api/citizens)', level=2)

add_endpoint_section(doc, 'GET', '/api/citizens', 'Retrieve all registered citizens.',
    res_example='[{"national_id": "NA-0001", "first_name": "Alem", "last_name": "Tesfaye", "date_of_birth": "1990-05-15", "gender": "Female", "phone": "+251911000001", "region": "Amhara", "district": "Bahir Dar", "village": "Kebele 01", "registered_by": "usr_001", "latitude": 11.59, "longitude": 37.39}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'GET', '/api/citizens/national/:nationalId', 'Retrieve a citizen by National ID.',
    params=[['nationalId', 'string', 'Yes', 'National ID number (path parameter)']],
    res_example='{"national_id": "NA-0001", "first_name": "Alem", "last_name": "Tesfaye", "date_of_birth": "1990-05-15", "gender": "Female", "id_type": "National ID", "id_number": "1234567890"}',
    errors=[['404', 'Citizen not found'], ['500', 'Server error']]
)

add_endpoint_section(doc, 'POST', '/api/citizens', 'Register a new citizen for National ID.',
    req_body=[
        ['nationalId', 'string', 'Yes', 'Unique National ID'],
        ['firstName', 'string', 'Yes', 'First name'],
        ['lastName', 'string', 'Yes', 'Last name'],
        ['grandfatherName', 'string', 'No', 'Grandfather name'],
        ['dateOfBirth', 'date', 'Yes', 'Date of birth (YYYY-MM-DD)'],
        ['gender', 'string', 'Yes', 'Gender: Male, Female, or Other'],
        ['phone', 'string', 'Yes', 'Phone number'],
        ['email', 'string', 'No', 'Email address'],
        ['address', 'string', 'No', 'Street address'],
        ['region', 'string', 'Yes', 'Region name'],
        ['district', 'string', 'Yes', 'District name'],
        ['village', 'string', 'Yes', 'Village name'],
        ['occupation', 'string', 'No', 'Occupation'],
        ['maritalStatus', 'string', 'No', 'Marital status'],
        ['registeredBy', 'string', 'Yes', 'Registering officer user ID'],
        ['registeredByName', 'string', 'Yes', 'Registering officer name'],
        ['idType', 'string', 'No', 'Document type: National ID, Birth Certificate, Passport'],
        ['idNumber', 'string', 'No', 'Document number'],
        ['biometrics', 'boolean', 'No', 'Biometrics collected (true/false)'],
        ['latitude', 'number', 'No', 'GPS latitude'],
        ['longitude', 'number', 'No', 'GPS longitude'],
    ],
    res_example='{"national_id": "NA-0001", "first_name": "Alem", "last_name": "Tesfaye", "date_of_birth": "1990-05-15", "gender": "Female", "region": "Amhara", "registered_by": "usr_001", "latitude": 11.59, "longitude": 37.39}',
    errors=[['409', 'Duplicate citizen (name + grandfather name already exists)'], ['500', 'Server error']]
)

# ---- REPORTS ----
doc.add_heading('Reports (/api/reports)', level=2)

add_endpoint_section(doc, 'GET', '/api/reports', 'Retrieve all field reports.',
    res_example='[{"id": 1, "report_id": "RPT-2026-001", "employee_id": "FO-001", "employee_name": "Abebe Kebede", "report_date": "2026-08-17", "site_name": "Site A", "registrations": 25, "operational_status": "active", "work_hours": 8.5, "submitted_at": "2026-08-17T17:00:00Z"}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'GET', '/api/reports/:id', 'Retrieve a single report by ID.',
    params=[['id', 'integer', 'Yes', 'Report ID (path parameter)']],
    res_example='{"id": 1, "report_id": "RPT-2026-001", "employee_id": "FO-001", "employee_name": "Abebe Kebede", "report_date": "2026-08-17", "site_name": "Site A", "registrations": 25, "activities": "Citizen registration drive", "latitude": 9.02, "longitude": 38.74}',
    errors=[['404', 'Report not found'], ['500', 'Server error']]
)

add_endpoint_section(doc, 'POST', '/api/reports', 'Submit a new field report.',
    req_body=[
        ['reportId', 'string', 'Yes', 'Unique report ID'],
        ['employeeId', 'string', 'Yes', 'Officer employee ID'],
        ['employeeName', 'string', 'Yes', 'Officer name'],
        ['reportDate', 'date', 'Yes', 'Report date'],
        ['region', 'string', 'Yes', 'Region'],
        ['siteName', 'string', 'Yes', 'Site name'],
        ['registrations', 'integer', 'No', 'Number of registrations completed'],
        ['operationalStatus', 'string', 'No', 'Operational status'],
        ['workHours', 'number', 'No', 'Hours worked'],
        ['activities', 'string', 'No', 'Activity description'],
        ['challenges', 'string', 'No', 'Challenges encountered'],
        ['latitude', 'number', 'No', 'GPS latitude'],
        ['longitude', 'number', 'No', 'GPS longitude'],
    ],
    res_example='{"id": 1, "report_id": "RPT-2026-001", "employee_id": "FO-001", "employee_name": "Abebe Kebede", "report_date": "2026-08-17", "site_name": "Site A", "registrations": 25, "submitted_at": "2026-08-17T17:00:00Z"}',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'PUT', '/api/reports/:id', 'Update an existing field report.',
    params=[['id', 'integer', 'Yes', 'Report ID (path parameter)']],
    req_body=[
        ['siteName', 'string', 'No', 'Updated site name'],
        ['registrations', 'integer', 'No', 'Updated registration count'],
        ['activities', 'string', 'No', 'Updated activities'],
        ['challenges', 'string', 'No', 'Updated challenges'],
        ['comments', 'string', 'No', 'Updated comments'],
    ],
    res_example='{"id": 1, "report_id": "RPT-2026-001", "site_name": "Site A Updated", "registrations": 30}',
    errors=[['404', 'Report not found'], ['500', 'Server error']]
)

add_endpoint_section(doc, 'DELETE', '/api/reports/:id', 'Delete a field report.',
    params=[['id', 'integer', 'Yes', 'Report ID (path parameter)']],
    res_example='{"message": "Report deleted successfully"}',
    errors=[['404', 'Report not found'], ['500', 'Server error']]
)

# ---- TASKS ----
doc.add_heading('Tasks (/api/tasks)', level=2)

add_endpoint_section(doc, 'GET', '/api/tasks', 'Retrieve all tasks.',
    res_example='[{"id": "tsk_001", "employee_id": "FO-001", "assigned_by": "SUP-001", "assigned_by_name": "Fatuma Ali", "title": "Register citizens in Kebele 05", "description": "Complete registration drive", "deadline": "2026-08-20", "priority": "high", "status": "pending", "created_at": "2026-08-17T09:00:00Z"}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'GET', '/api/tasks/employee/:employeeId', 'Retrieve tasks assigned to a specific employee.',
    params=[['employeeId', 'string', 'Yes', 'Employee ID (path parameter)']],
    res_example='[{"id": "tsk_001", "title": "Register citizens in Kebele 05", "priority": "high", "status": "pending", "deadline": "2026-08-20"}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'POST', '/api/tasks', 'Create and assign a new task.',
    req_body=[
        ['id', 'string', 'Yes', 'Unique task ID'],
        ['employeeId', 'string', 'Yes', 'Assigned officer employee ID'],
        ['assignedBy', 'string', 'Yes', 'Supervisor employee ID'],
        ['assignedByName', 'string', 'Yes', 'Supervisor name'],
        ['title', 'string', 'Yes', 'Task title'],
        ['description', 'string', 'No', 'Task description'],
        ['deadline', 'date', 'Yes', 'Due date'],
        ['priority', 'string', 'Yes', 'Priority: low, medium, high'],
    ],
    res_example='{"id": "tsk_001", "employee_id": "FO-001", "title": "Register citizens in Kebele 05", "priority": "high", "status": "pending", "created_at": "2026-08-17T09:00:00Z"}',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'PUT', '/api/tasks/:id', 'Update a task (status, priority, details).',
    params=[['id', 'string', 'Yes', 'Task ID (path parameter)']],
    req_body=[
        ['status', 'string', 'No', 'Updated status: pending, in-progress, completed'],
        ['priority', 'string', 'No', 'Updated priority'],
        ['title', 'string', 'No', 'Updated title'],
        ['description', 'string', 'No', 'Updated description'],
    ],
    res_example='{"id": "tsk_001", "title": "Register citizens in Kebele 05", "status": "completed", "completed_at": "2026-08-19T14:30:00Z"}',
    errors=[['404', 'Task not found'], ['500', 'Server error']]
)

add_endpoint_section(doc, 'DELETE', '/api/tasks/:id', 'Delete a task.',
    params=[['id', 'string', 'Yes', 'Task ID (path parameter)']],
    res_example='{"message": "Task deleted successfully"}',
    errors=[['404', 'Task not found'], ['500', 'Server error']]
)

# ---- PERMISSIONS ----
doc.add_heading('Permissions (/api/permissions)', level=2)

add_endpoint_section(doc, 'GET', '/api/permissions', 'Retrieve all permission requests.',
    res_example='[{"id": "perm_001", "employee_id": "FO-001", "employee_name": "Abebe Kebede", "permission_type": "leave", "start_date": "2026-08-20", "end_date": "2026-08-22", "reason": "Personal leave", "status": "pending", "requested_at": "2026-08-17T10:00:00Z"}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'POST', '/api/permissions', 'Submit a new permission request.',
    req_body=[
        ['id', 'string', 'Yes', 'Unique permission ID'],
        ['employeeId', 'string', 'Yes', 'Requesting officer employee ID'],
        ['employeeName', 'string', 'Yes', 'Requesting officer name'],
        ['permissionType', 'string', 'Yes', 'Type: leave, travel, equipment, other'],
        ['startDate', 'date', 'Yes', 'Start date'],
        ['endDate', 'date', 'Yes', 'End date'],
        ['reason', 'string', 'Yes', 'Reason for request'],
    ],
    res_example='{"id": "perm_001", "employee_id": "FO-001", "permission_type": "leave", "status": "pending", "requested_at": "2026-08-17T10:00:00Z"}',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'PUT', '/api/permissions/:id', 'Approve or reject a permission request.',
    params=[['id', 'string', 'Yes', 'Permission ID (path parameter)']],
    req_body=[
        ['status', 'string', 'Yes', 'Updated status: approved or rejected'],
        ['approvedBy', 'string', 'Yes', 'Approving supervisor employee ID'],
        ['rejectReason', 'string', 'No', 'Rejection reason (if status is rejected)'],
    ],
    res_example='{"id": "perm_001", "status": "approved", "approved_by": "SUP-001", "approved_at": "2026-08-17T11:00:00Z"}',
    errors=[['404', 'Permission not found'], ['500', 'Server error']]
)

# ---- ALERTS ----
doc.add_heading('Alerts (/api/alerts)', level=2)

add_endpoint_section(doc, 'GET', '/api/alerts', 'Retrieve all alerts and messages.',
    res_example='[{"id": "alt_001", "title": "Task Assigned", "message": "You have been assigned a new task", "priority": "medium", "type": "task", "read": false, "sent_by": "SUP-001", "sent_by_name": "Fatuma Ali", "timestamp": "2026-08-17T09:00:00Z"}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'POST', '/api/alerts', 'Send a new alert or message.',
    req_body=[
        ['id', 'string', 'Yes', 'Unique alert ID'],
        ['title', 'string', 'Yes', 'Alert title'],
        ['message', 'string', 'Yes', 'Alert message body'],
        ['priority', 'string', 'No', 'Priority: low, medium, high (default: medium)'],
        ['type', 'string', 'No', 'Alert type (e.g., emergency, task, general)'],
        ['targetAll', 'boolean', 'No', 'Send to all users (default: false)'],
        ['targetEmployeeId', 'string', 'No', 'Specific target employee ID'],
        ['sentBy', 'string', 'Yes', 'Sender employee ID'],
        ['sentByName', 'string', 'Yes', 'Sender name'],
        ['sentByRole', 'string', 'No', 'Sender role'],
    ],
    res_example='{"id": "alt_001", "title": "Task Assigned", "message": "You have been assigned a new task", "priority": "medium", "sent_by": "SUP-001", "timestamp": "2026-08-17T09:00:00Z"}',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'PUT', '/api/alerts/:id', 'Mark an alert as read or unread.',
    params=[['id', 'string', 'Yes', 'Alert ID (path parameter)']],
    req_body=[['read', 'boolean', 'No', 'Read status (defaults to true)']],
    res_example='{"id": "alt_001", "read": true}',
    errors=[['404', 'Alert not found'], ['500', 'Server error']]
)

add_endpoint_section(doc, 'DELETE', '/api/alerts', 'Clear all alerts.',
    res_example='{"message": "All alerts cleared"}',
    errors=[['500', 'Server error']]
)

# ---- AUDIT ----
doc.add_heading('Audit Logs (/api/audit)', level=2)

add_endpoint_section(doc, 'GET', '/api/audit', 'Retrieve all audit log entries.',
    res_example='[{"id": "aud_001", "user_id": "usr_001", "user_name": "Abebe Kebede", "action": "citizen_registered", "details": "Registered citizen NA-0001", "timestamp": "2026-08-17T10:30:00Z", "ip": "192.168.1.100"}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'POST', '/api/audit', 'Create a new audit log entry.',
    req_body=[
        ['id', 'string', 'Yes', 'Unique audit ID'],
        ['userId', 'string', 'Yes', 'User who performed the action'],
        ['userName', 'string', 'Yes', 'User name'],
        ['action', 'string', 'Yes', 'Action performed (e.g., citizen_registered, report_submitted)'],
        ['details', 'string', 'No', 'Additional details about the action'],
        ['ip', 'string', 'No', 'Client IP address'],
    ],
    res_example='{"id": "aud_001", "user_id": "usr_001", "action": "citizen_registered", "details": "Registered citizen NA-0001", "timestamp": "2026-08-17T10:30:00Z"}',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'DELETE', '/api/audit/:id', 'Delete a specific audit log entry.',
    params=[['id', 'string', 'Yes', 'Audit log ID (path parameter)']],
    res_example='{"message": "Audit record deleted"}',
    errors=[['404', 'Audit record not found'], ['500', 'Server error']]
)

add_endpoint_section(doc, 'DELETE', '/api/audit', 'Clear all audit logs.',
    res_example='{"message": "All audit logs cleared"}',
    errors=[['500', 'Server error']]
)

# ---- SCREEN TIME ----
doc.add_heading('Screen Time (/api/screentime)', level=2)

add_endpoint_section(doc, 'GET', '/api/screentime', 'Retrieve all screen time records.',
    res_example='[{"id": "scr_001", "employee_id": "FO-001", "employee_name": "Abebe Kebede", "date": "2026-08-17", "login_time": "08:00", "logout_time": "17:30", "total_screen_time": 30600, "trust_score": 95}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'GET', '/api/screentime/employee/:employeeId', 'Retrieve screen time records for a specific employee.',
    params=[['employeeId', 'string', 'Yes', 'Employee ID (path parameter)']],
    res_example='[{"id": "scr_001", "date": "2026-08-17", "login_time": "08:00", "total_screen_time": 30600, "trust_score": 95}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'POST', '/api/screentime', 'Create or update a screen time record.',
    req_body=[
        ['id', 'string', 'Yes', 'Unique screen time record ID'],
        ['employeeId', 'string', 'Yes', 'Employee ID'],
        ['employeeName', 'string', 'Yes', 'Employee name'],
        ['date', 'date', 'Yes', 'Record date'],
        ['loginTime', 'string', 'No', 'Login time (HH:MM)'],
        ['logoutTime', 'string', 'No', 'Logout time (HH:MM)'],
        ['totalScreenTime', 'integer', 'No', 'Total active screen time (seconds)'],
        ['idleTime', 'integer', 'No', 'Idle time (seconds)'],
        ['trustScore', 'integer', 'No', 'Trust score (0-100)'],
        ['isLoggedIn', 'boolean', 'No', 'Currently logged in'],
    ],
    res_example='{"id": "scr_001", "employee_id": "FO-001", "date": "2026-08-17", "total_screen_time": 30600, "trust_score": 95}',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'PUT', '/api/screentime/:id', 'Update screen time settings or verification status.',
    params=[['id', 'string', 'Yes', 'Screen time record ID (path parameter)']],
    req_body=[
        ['screenTimeLimit', 'integer', 'No', 'Updated daily limit (seconds)'],
        ['verified', 'boolean', 'No', 'Verification status'],
        ['verifiedBy', 'string', 'No', 'Verifier employee ID'],
    ],
    res_example='{"id": "scr_001", "screen_time_limit": 28800, "verified": true, "verified_by": "SUP-001"}',
    errors=[['404', 'Record not found'], ['500', 'Server error']]
)

add_endpoint_section(doc, 'DELETE', '/api/screentime/:id', 'Delete a screen time record.',
    params=[['id', 'string', 'Yes', 'Record ID (path parameter)']],
    res_example='{"message": "Screen time record deleted"}',
    errors=[['404', 'Record not found'], ['500', 'Server error']]
)

# ---- VERIFICATION ----
doc.add_heading('Verification (/api/verification)', level=2)

add_endpoint_section(doc, 'GET', '/api/verification', 'Retrieve all verification history records.',
    res_example='[{"id": "vrf_001", "officer_id": "FO-001", "officer_name": "Abebe Kebede", "question": "What is your current location?", "answer": "Kebele 05, Bahir Dar", "success": true, "score": 100, "response_time": 12, "timestamp": "2026-08-17T10:15:00Z"}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'GET', '/api/verification/officer/:officerId', 'Retrieve verification records for a specific officer.',
    params=[['officerId', 'string', 'Yes', 'Officer employee ID (path parameter)']],
    res_example='[{"id": "vrf_001", "question": "What is your current location?", "success": true, "score": 100, "response_time": 12}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'POST', '/api/verification', 'Record a verification response.',
    req_body=[
        ['id', 'string', 'Yes', 'Unique verification ID'],
        ['officerId', 'string', 'Yes', 'Officer employee ID'],
        ['officerName', 'string', 'Yes', 'Officer name'],
        ['question', 'string', 'Yes', 'Verification question'],
        ['answer', 'string', 'Yes', 'Officer answer'],
        ['success', 'boolean', 'No', 'Whether verification passed'],
        ['score', 'integer', 'No', 'Score (0-100)'],
        ['responseTime', 'integer', 'No', 'Response time in seconds'],
        ['message', 'string', 'No', 'System message'],
        ['penalties', 'array', 'No', 'Array of penalty objects'],
    ],
    res_example='{"id": "vrf_001", "officer_id": "FO-001", "question": "What is your current location?", "success": true, "score": 100, "response_time": 12, "timestamp": "2026-08-17T10:15:00Z"}',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'DELETE', '/api/verification/:id', 'Delete a specific verification record.',
    params=[['id', 'string', 'Yes', 'Verification record ID (path parameter)']],
    res_example='{"message": "Verification record deleted"}',
    errors=[['404', 'Record not found'], ['500', 'Server error']]
)

add_endpoint_section(doc, 'DELETE', '/api/verification', 'Clear all verification records.',
    res_example='{"message": "All verification records cleared"}',
    errors=[['500', 'Server error']]
)

# ---- LOCATIONS ----
doc.add_heading('Locations (/api/locations)', level=2)

add_endpoint_section(doc, 'GET', '/api/locations/level/:level', 'Retrieve all locations at a specific level.',
    params=[['level', 'string', 'Yes', 'Location level: country, region, zone, woreda, kebele']],
    res_example='[{"id": 1, "name": "Amhara", "parent_id": 1}, {"id": 2, "name": "Oromia", "parent_id": 1}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'GET', '/api/locations/children/:parentId', 'Retrieve child locations of a parent.',
    params=[['parentId', 'integer', 'Yes', 'Parent location ID (path parameter)']],
    res_example='[{"id": 5, "name": "Bahir Dar", "level": "zone"}, {"id": 6, "name": "Gondar", "level": "zone"}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'GET', '/api/locations/communities', 'Retrieve communities within a kebele.',
    params=[['kebele_id', 'integer', 'Yes', 'Kebele ID (query parameter)']],
    res_example='[{"id": 1, "name": "Community A"}, {"id": 2, "name": "Community B"}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'GET', '/api/locations/:id', 'Retrieve a single location by ID.',
    params=[['id', 'integer', 'Yes', 'Location ID (path parameter)']],
    res_example='{"id": 1, "name": "Amhara", "level": "region", "parent_id": 1}',
    errors=[['404', 'Location not found'], ['500', 'Server error']]
)

# ---- SUPERVISOR REPORTS ----
doc.add_heading('Supervisor Reports (/api/supervisor-reports)', level=2)

add_endpoint_section(doc, 'GET', '/api/supervisor-reports', 'Retrieve all supervisor reports.',
    res_example='[{"id": "srp_001", "supervisor_id": "SUP-001", "supervisor_name": "Fatuma Ali", "officer_id": "FO-001", "officer_name": "Abebe Kebede", "report_date": "2026-08-17", "performance": "excellent", "overall_rating": 9, "status": "submitted", "type": "officer_report"}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'GET', '/api/supervisor-reports/supervisor/:supervisorId', 'Retrieve reports created by a specific supervisor.',
    params=[['supervisorId', 'string', 'Yes', 'Supervisor employee ID (path parameter)']],
    res_example='[{"id": "srp_001", "officer_name": "Abebe Kebede", "performance": "excellent", "overall_rating": 9}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'GET', '/api/supervisor-reports/officer/:officerId', 'Retrieve reports about a specific officer.',
    params=[['officerId', 'string', 'Yes', 'Officer employee ID (path parameter)']],
    res_example='[{"id": "srp_001", "supervisor_name": "Fatuma Ali", "performance": "excellent", "overall_rating": 9}]',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'POST', '/api/supervisor-reports', 'Create a new supervisor report (officer report or self report).',
    req_body=[
        ['id', 'string', 'Yes', 'Unique report ID'],
        ['supervisorId', 'string', 'Yes', 'Supervisor employee ID'],
        ['supervisorName', 'string', 'Yes', 'Supervisor name'],
        ['officerId', 'string', 'No', 'Officer employee ID (for officer_report type)'],
        ['officerName', 'string', 'No', 'Officer name (for officer_report type)'],
        ['reportDate', 'date', 'Yes', 'Report date'],
        ['type', 'string', 'Yes', 'Report type: officer_report or self_report'],
        ['performance', 'string', 'No', 'Performance rating: excellent, good, satisfactory, poor'],
        ['attendance', 'string', 'No', 'Attendance assessment'],
        ['quality', 'string', 'No', 'Quality assessment'],
        ['punctuality', 'string', 'No', 'Punctuality assessment'],
        ['teamwork', 'string', 'No', 'Teamwork assessment'],
        ['communication', 'string', 'No', 'Communication assessment'],
        ['overallRating', 'integer', 'No', 'Overall rating (1-10)'],
        ['comments', 'string', 'No', 'Supervisor comments'],
        ['recommendations', 'string', 'No', 'Recommendations'],
    ],
    res_example='{"id": "srp_001", "supervisor_id": "SUP-001", "officer_id": "FO-001", "type": "officer_report", "overall_rating": 9, "status": "submitted"}',
    errors=[['500', 'Server error']]
)

add_endpoint_section(doc, 'PUT', '/api/supervisor-reports/:id', 'Update a supervisor report.',
    params=[['id', 'string', 'Yes', 'Report ID (path parameter)']],
    req_body=[
        ['performance', 'string', 'No', 'Updated performance rating'],
        ['overallRating', 'integer', 'No', 'Updated overall rating'],
        ['comments', 'string', 'No', 'Updated comments'],
        ['status', 'string', 'No', 'Updated status'],
    ],
    res_example='{"id": "srp_001", "overall_rating": 10, "status": "approved"}',
    errors=[['404', 'Report not found'], ['500', 'Server error']]
)

add_endpoint_section(doc, 'DELETE', '/api/supervisor-reports/:id', 'Delete a supervisor report.',
    params=[['id', 'string', 'Yes', 'Report ID (path parameter)']],
    res_example='{"message": "Report deleted"}',
    errors=[['404', 'Report not found'], ['500', 'Server error']]
)

# ---- SYNC ----
doc.add_heading('Sync (/api/sync)', level=2)

add_endpoint_section(doc, 'POST', '/api/sync', 'Synchronize offline data to the server. The type field determines the data model and operation.',
    req_body=[
        ['type', 'string', 'Yes', 'Data type: report, citizen, permission, task, user, alert, audit, screen_time, verification, supervisor_report (and _update/_delete variants)'],
        ['data', 'object', 'Yes', 'Data payload matching the specified type schema'],
    ],
    req_example='POST /api/sync\nContent-Type: application/json\nAuthorization: Bearer <token>\n\n{\n  "type": "citizen",\n  "data": {\n    "nationalId": "NA-0001",\n    "firstName": "Alem",\n    "lastName": "Tesfaye",\n    "dateOfBirth": "1990-05-15",\n    "gender": "Female",\n    "phone": "+251911000001",\n    "region": "Amhara",\n    "district": "Bahir Dar",\n    "village": "Kebele 01",\n    "registeredBy": "usr_001",\n    "registeredByName": "Abebe Kebede",\n    "latitude": 11.59,\n    "longitude": 37.39\n  }\n}',
    res_example='{"success": true, "data": {"national_id": "NA-0001", "first_name": "Alem", "last_name": "Tesfaye"}, "message": "Citizen synced successfully"}',
    errors=[['409', 'Duplicate citizen (for citizen type)'], ['500', 'Server error or invalid type']]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Supported Sync Types:')
run.bold = True
make_table(doc, ['Type', 'Operation', 'Description'], [
    ['report', 'Upsert', 'Sync a field report (upserts on report_id)'],
    ['citizen', 'Upsert', 'Sync a citizen registration (upserts on national_id)'],
    ['permission', 'Upsert', 'Sync a permission request'],
    ['permission_update', 'Update', 'Update permission approval status'],
    ['task', 'Upsert', 'Sync a task assignment'],
    ['task_update', 'Update', 'Update task completion status'],
    ['user', 'Upsert', 'Sync a user account'],
    ['user_update', 'Update', 'Update user profile'],
    ['user_status_update', 'Update', 'Update user active/inactive status'],
    ['user_delete', 'Delete', 'Delete a user account'],
    ['alert', 'Upsert', 'Sync an alert message'],
    ['alert_read', 'Update', 'Mark alert as read'],
    ['audit', 'Upsert', 'Sync an audit log entry'],
    ['screen_time', 'Upsert', 'Sync a screen time record'],
    ['screen_time_update', 'Update', 'Update screen time settings'],
    ['screen_time_delete', 'Delete', 'Delete a screen time record'],
    ['verification', 'Upsert', 'Sync a verification record'],
    ['verification_delete', 'Delete', 'Delete a verification record'],
    ['supervisor_report', 'Upsert', 'Sync a supervisor report'],
])

doc.add_page_break()

# =============================================
# CHANGE LOG
# =============================================
doc.add_heading('Change Log', level=1)
make_table(doc, ['Date', 'Change'], [
    [datetime.date.today().strftime('%Y-%m-%d'), 'Initial API documentation release (v1.0.0). Documented all 47 endpoints across 13 resource modules.'],
])

doc.add_paragraph()

# =============================================
# GLOSSARY
# =============================================
doc.add_heading('Glossary / Definitions', level=1)
make_table(doc, ['Term', 'Definition'], [
    ['API', 'Application Programming Interface -- a set of protocols for building and interacting with software applications.'],
    ['bcrypt', 'A password hashing algorithm designed to resist brute-force attacks through computational cost.'],
    ['CRUD', 'Create, Read, Update, Delete -- the four fundamental database operations.'],
    ['CORS', 'Cross-Origin Resource Sharing -- a security mechanism controlling cross-domain HTTP requests.'],
    ['Dexie.js', 'A JavaScript library providing a simplified wrapper around the browser IndexedDB API.'],
    ['JWT', 'JSON Web Token -- a compact token format for securely transmitting information between parties.'],
    ['MVC', 'Model-View-Controller -- an architectural pattern separating data, presentation, and logic.'],
    ['Multer', 'A Node.js middleware for handling file uploads via multipart/form-data.'],
    ['Nodemailer', 'A Node.js library for sending emails via SMTP.'],
    ['PWA', 'Progressive Web App -- a web application providing native app-like capabilities including offline support.'],
    ['RBAC', 'Role-Based Access Control -- access regulation based on assigned user roles.'],
    ['REST', 'Representational State Transfer -- an architectural style for networked applications.'],
    ['SPA', 'Single Page Application -- a web app that loads once and dynamically updates content.'],
    ['Sync Service', 'A client-side module managing offline data queuing and server synchronization.'],
    ['Upsert', 'An operation that inserts a new record if it does not exist, or updates it if it does.'],
    ['Zod', 'A TypeScript-first schema validation library for form and API input validation.'],
])

# =============================================
# SAVE
# =============================================
output_path = 'C:\\Users\\nebi\\Desktop\\mongoreact\\fieldsync\\FieldSync_API_Documentation_v2.docx'
doc.save(output_path)
print(f'API Documentation saved to: {output_path}')
