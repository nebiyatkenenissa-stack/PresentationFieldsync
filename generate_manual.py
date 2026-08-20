from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# --- Style Configuration ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(30, 58, 95)
    hs.font.name = 'Calibri'

# --- Helper Functions ---
def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells_data, bold=False, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        if bold or header:
            run.bold = True
        if header:
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_shading(cell, '1E3A5F')
    return row

def make_table(doc, headers, rows):
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, headers, header=True)
    for row_data in rows:
        add_table_row(table, row_data)
    return table

# =============================================
# COVER PAGE
# =============================================
for _ in range(6):
    doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('FieldSync')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(30, 58, 95)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('User Manual')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(79, 195, 247)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Comprehensive Guide to Installation, Usage, and Maintenance')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f'Version 1.0.0\n{datetime.date.today().strftime("%B %Y")}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 116, 139)

doc.add_page_break()

# =============================================
# TABLE OF CONTENTS
# =============================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    ('1.', 'Introduction'),
    ('2.', 'System Requirements'),
    ('3.', 'Installation Instructions'),
    ('4.', 'Getting Started'),
    ('5.', 'Features Overview'),
    ('6.', 'Usage Instructions'),
    ('7.', 'Troubleshooting Guide'),
    ('8.', 'Frequently Asked Questions (FAQ)'),
    ('9.', 'Support and Contact'),
    ('10.', 'Revision History'),
]
for num, item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  {item}')
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(4)

doc.add_page_break()

# =============================================
# 1. INTRODUCTION
# =============================================
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'FieldSync is a comprehensive field workforce management and citizen registration '
    'platform designed to streamline the operations of government field officers, supervisors, '
    'and managers. The system enables real-time citizen registration for National ID programs, '
    'field reporting, task assignment and tracking, team management, and performance analytics '
    '--- all with full offline capability and automatic data synchronization when connectivity '
    'is restored.'
)
doc.add_paragraph(
    'This manual aims to guide users through installation, usage, and maintenance of the system. '
    'It covers all three user roles: Field Officers (who register citizens and submit reports), '
    'Supervisors (who manage teams and oversee operations), and Managers (who have full '
    'administrative access including analytics and audit logs).'
)

# =============================================
# 2. SYSTEM REQUIREMENTS
# =============================================
doc.add_heading('2. System Requirements', level=1)
doc.add_paragraph('Before installing FieldSync, ensure your environment meets the following requirements:')

doc.add_heading('2.1 Server Requirements', level=2)
server_reqs = [
    ['Operating System', 'Linux (Ubuntu 20.04+), macOS 12+, or Windows 10+ with Docker Desktop'],
    ['Docker', 'Docker Engine 20.10+ and Docker Compose v2+'],
    ['RAM', 'Minimum 4 GB (8 GB recommended for production)'],
    ['Disk Space', 'Minimum 10 GB free space'],
    ['CPU', '2 or more cores recommended'],
    ['Network', 'Port 5001 (backend) and port 30001 (frontend) must be accessible'],
]
make_table(doc, ['Requirement', 'Description'], server_reqs)

doc.add_paragraph()
doc.add_heading('2.2 Database Requirements', level=2)
db_reqs = [
    ['PostgreSQL', 'Version 16 (Alpine image used via Docker)'],
    ['Database Name', 'fieldsync_db (configurable via environment variables)'],
    ['Connection', 'Default port 5432, configured automatically via Docker networking'],
]
make_table(doc, ['Requirement', 'Description'], db_reqs)

doc.add_paragraph()
doc.add_heading('2.3 Client / Browser Requirements', level=2)
client_reqs = [
    ['Web Browser', 'Chrome 90+, Firefox 88+, Edge 90+, or Safari 14+'],
    ['Screen Resolution', 'Minimum 1024 x 768 (responsive design supports mobile)'],
    ['Internet Connection', 'Required for initial setup; offline mode available after first login'],
    ['GPS', 'Required for location-based features (field officers)'],
    ['Camera', 'Optional, for profile photos and citizen document capture'],
]
make_table(doc, ['Requirement', 'Description'], client_reqs)

doc.add_paragraph()
doc.add_heading('2.4 Environment Variables', level=2)
doc.add_paragraph('The following environment variables must be configured before deployment:')
env_vars = [
    ['PORT', '5000', 'Backend server port'],
    ['DB_USER', 'postgres', 'PostgreSQL username'],
    ['DB_PASSWORD', '(set securely)', 'PostgreSQL password'],
    ['DB_HOST', 'db', 'Database host (use "db" in Docker)'],
    ['DB_PORT', '5432', 'PostgreSQL port'],
    ['DB_NAME', 'fieldsync_db', 'Database name'],
    ['JWT_SECRET', '(set securely)', 'Secret key for JWT token generation'],
    ['EMAIL_USER', '(your email)', 'Email address for notifications'],
    ['EMAIL_PASS', '(app password)', 'Email application password'],
]
make_table(doc, ['Variable', 'Default', 'Description'], env_vars)

# =============================================
# 3. INSTALLATION INSTRUCTIONS
# =============================================
doc.add_heading('3. Installation Instructions', level=1)
doc.add_paragraph(
    'FieldSync is deployed using Docker Compose, which sets up all required services '
    '(database, backend API, and frontend) automatically.'
)

doc.add_heading('3.1 Step-by-Step Installation', level=2)

steps = [
    ('Download the source code', 'Obtain the FieldSync package from your system administrator or the official repository. Extract the archive to a location of your choice.'),
    ('Configure environment variables', 'Copy the .env.example file to .env in the project root directory. Open the .env file and fill in your database credentials, JWT secret, and email configuration.'),
    ('Start Docker services', 'Open a terminal in the project root directory and run the command: docker compose up -d --build. This will build the images and start all containers.'),
    ('Verify installation', 'Open your browser and navigate to http://localhost:30001. You should see the FieldSync login page. The backend API is accessible at http://localhost:5001.'),
]
for i, (title, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {i}: {title}')
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph(desc)

doc.add_paragraph()
doc.add_heading('3.2 Docker Services Overview', level=2)
services = [
    ['fieldsync-db', 'PostgreSQL 16 Alpine', 'Database', 'Internal (5432)'],
    ['fieldsync-backend', 'Node.js/Express + TypeScript', 'REST API', '5001:5000'],
    ['fieldsync-frontend', 'React + Vite + Nginx', 'Web Application', '30001:80'],
]
make_table(doc, ['Container', 'Technology', 'Role', 'Port Mapping'], services)

doc.add_paragraph()
doc.add_heading('3.3 Installation Notes', level=2)
doc.add_paragraph(
    'Docker and Docker Compose must be installed and running on your system before starting '
    'the installation. On Windows, install Docker Desktop. On Linux, install docker.io and '
    'docker-compose-plugin.',
    style='List Bullet'
)
doc.add_paragraph(
    'Ensure ports 5001 and 30001 are not already in use by other applications.',
    style='List Bullet'
)
doc.add_paragraph(
    'The database data is persisted in a Docker volume named "pgdata". Uploaded files are '
    'persisted in a volume named "uploads".',
    style='List Bullet'
)
doc.add_paragraph(
    'For production deployments, change all default passwords and the JWT secret to strong, '
    'unique values.',
    style='List Bullet'
)

# =============================================
# 4. GETTING STARTED
# =============================================
doc.add_heading('4. Getting Started', level=1)
doc.add_paragraph(
    'This section describes how to start using FieldSync after installation.'
)

doc.add_heading('4.1 Launching the Application', level=2)
doc.add_paragraph('Open your web browser and navigate to the application URL (e.g., http://localhost:30001). You will see the FieldSync login screen with the brand logo and feature highlights.')

doc.add_heading('4.2 Logging In', level=2)
login_steps = [
    'Enter your registered email address in the Email field.',
    'Enter your password in the Password field. Use the eye icon to toggle password visibility.',
    'Click the "Sign In" button to authenticate.',
    'The system will verify your credentials and redirect you to the dashboard based on your assigned role.',
]
for step in login_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('4.3 User Roles', level=2)
doc.add_paragraph('FieldSync provides three distinct user roles, each with specific access levels and capabilities:')
roles = [
    ['Field Officer', 'Register citizens for National ID, submit field reports, manage permissions, view alerts'],
    ['Supervisor', 'Manage team members, assign and track tasks, view supervisor reports, monitor screen time, verify officer activities'],
    ['Manager', 'Full administrative access, manage all users, view analytics and audit logs, access all reports and citizen database'],
]
make_table(doc, ['Role', 'Primary Capabilities'], roles)

doc.add_paragraph()
doc.add_heading('4.4 Initial Setup', level=2)
doc.add_paragraph(
    'After your first login, follow these recommended steps:',
)
init_steps = [
    'Navigate to your Profile page to review and update your personal information.',
    'Upload a profile photo if one has not been set.',
    'Verify that your assigned role and team information are correct.',
    'Familiarize yourself with the sidebar navigation, which adapts based on your role.',
    'Supervisors and Managers: Create user accounts for your team members via the User Management section.',
]
for step in init_steps:
    doc.add_paragraph(step, style='List Number')

# =============================================
# 5. FEATURES OVERVIEW
# =============================================
doc.add_heading('5. Features Overview', level=1)

doc.add_heading('5.1 Key Features', level=2)
features = [
    ['Citizen Registration', 'Register citizens for National ID with biographic data, document types (National ID, Birth Certificate, Passport), and GPS location capture.'],
    ['Field Reporting', 'Create, submit, and manage field reports with offline support. Reports are synced automatically when connectivity is restored.'],
    ['Task Management', 'Supervisors can assign tasks to field officers, set deadlines, and track completion status in real time.'],
    ['Team Management', 'Supervisors can view and manage team members, track performance, and monitor field activities.'],
    ['Permission Management', 'Officers can request permissions (travel, equipment, etc.) which flow through an approval workflow to supervisors.'],
    ['Supervisor Reports', 'Generate and view detailed supervisor reports summarizing team activities and performance metrics.'],
    ['Screen Time Tracking', 'Monitor application usage and screen time for field devices.'],
    ['Verification System', 'Periodic identity verification checks for field officers to confirm active duty status.'],
    ['Analytics Dashboard', 'Managers can access trend charts, registration statistics, and performance analytics.'],
    ['Audit Logs', 'Comprehensive audit trail of all system actions for accountability and compliance.'],
    ['Alerts & Messaging', 'Internal messaging system for notifications, alerts, and team communications.'],
    ['Offline Mode', 'Full offline capability with IndexedDB (Dexie) local storage and automatic sync when online.'],
    ['Multi-Language Support', 'Interface available in English, Amharic, Tigrinya, and Oromo with language auto-detection.'],
    ['GPS & Location', 'GPS capture for citizen registration and field officer location tracking.'],
    ['Dark/Light Theme', 'Toggle between dark and light themes for user preference and field conditions.'],
    ['Progressive Web App', 'Installable as a PWA for fast access and offline functionality.'],
]
make_table(doc, ['Feature', 'Description'], features)

doc.add_paragraph()
doc.add_heading('5.2 Offline Capability', level=2)
doc.add_paragraph(
    'FieldSync is designed for field use where internet connectivity may be unreliable. '
    'The application uses IndexedDB (via Dexie.js) to store data locally on the device. '
    'When the device goes offline, users can continue to register citizens and create reports. '
    'Once connectivity is restored, the SyncService automatically uploads all pending data to '
    'the server and resolves any conflicts.'
)
doc.add_paragraph(
    'A network status indicator is always visible in the interface, showing whether you are '
    'currently online or offline. Pending sync items are displayed with a badge counter.'
)

# =============================================
# 6. USAGE INSTRUCTIONS
# =============================================
doc.add_heading('6. Usage Instructions', level=1)
doc.add_paragraph('Detailed usage steps for performing common tasks:')

# --- 6.1 Registering a Citizen ---
doc.add_heading('6.1 Registering a Citizen (Field Officer)', level=2)
doc.add_paragraph(
    'Navigate to the "Register Citizen" section from the sidebar menu.',
)
doc.add_paragraph(
    'Fill in all required fields: First Name, Last Name, Date of Birth, Gender, Phone Number, and Region/District/Village.',
)
doc.add_paragraph(
    'Select the ID type (National ID, Birth Certificate, or Passport) and enter the ID number if available.',
)
doc.add_paragraph(
    'Toggle "Biometrics Collected" if biometric data has been captured.',
)
doc.add_paragraph(
    'Click "Register Citizen" to submit. GPS coordinates are captured automatically. The registration will be queued for sync if offline.',
)

# --- 6.2 Creating a Field Report ---
doc.add_heading('6.2 Creating a Field Report (Field Officer)', level=2)
doc.add_paragraph(
    'Navigate to "New Report" from the sidebar menu.',
)
doc.add_paragraph(
    'Enter the report title and detailed description of your field activity or observation.',
)
doc.add_paragraph(
    'Attach any supporting photos or documents using the upload feature.',
)
doc.add_paragraph(
    'Submit the report. It will be saved locally if offline and synced automatically when online.',
)

# --- 6.3 Assigning Tasks (Supervisor) ---
doc.add_heading('6.3 Assigning Tasks (Supervisor)', level=2)
doc.add_paragraph(
    'Navigate to "Tasks" from the sidebar menu.',
)
doc.add_paragraph(
    'Click to create a new task, enter the task title, description, and assign it to one or more field officers.',
)
doc.add_paragraph(
    'Set the priority level and due date for the task.',
)
doc.add_paragraph(
    'Save the task. Assigned officers will receive an alert notification.',
)

# --- 6.4 Managing Users (Manager) ---
doc.add_heading('6.4 Managing Users (Manager)', level=2)
doc.add_paragraph(
    'Navigate to "User Management" from the sidebar menu.',
)
doc.add_paragraph(
    'View the list of all registered users with their roles, status, and team assignments.',
)
doc.add_paragraph(
    'Click "Add User" to create a new account, specifying name, email, role (Manager/Supervisor/Field Officer), and team.',
)
doc.add_paragraph(
    'Edit or deactivate existing user accounts as needed.',
)

# --- 6.5 Viewing Analytics (Manager) ---
doc.add_heading('6.5 Viewing Analytics (Manager)', level=2)
doc.add_paragraph(
    'Navigate to "Analytics" from the sidebar menu.',
)
doc.add_paragraph(
    'View the registration trend charts and summary statistics.',
)
doc.add_paragraph(
    'Use filters to narrow down analytics by date range, region, or team.',
)
doc.add_paragraph(
    'Export or print reports for official use.',
)

# --- 6.6 Permission Requests (Field Officer) ---
doc.add_heading('6.6 Requesting Permissions (Field Officer)', level=2)
doc.add_paragraph(
    'Navigate to "Permissions" from the sidebar menu.',
)
doc.add_paragraph(
    'Click to create a new permission request and select the type (e.g., leave, travel, equipment).',
)
doc.add_paragraph(
    'Provide a reason and any required dates or details.',
)
doc.add_paragraph(
    'Submit the request. Your supervisor will be notified and can approve or reject it.',
)

# --- 6.7 Responding to Verification (Field Officer / Supervisor) ---
doc.add_heading('6.7 Identity Verification', level=2)
doc.add_paragraph(
    'Periodic verification prompts will appear on your screen requiring confirmation that you are the active user.',
)
doc.add_paragraph(
    'Click the confirmation button within the allotted time to verify your identity.',
)
doc.add_paragraph(
    'If verification is missed, a notification is sent to your supervisor for follow-up.',
)

# --- 6.8 Changing Language ---
doc.add_heading('6.8 Changing Language', level=2)
doc.add_paragraph(
    'Click the language selector in the sidebar or header to switch between English, Amharic, Tigrinya, and Oromo.',
)
doc.add_paragraph(
    'Your language preference is saved and persists across sessions.',
)

# =============================================
# 7. TROUBLESHOOTING GUIDE
# =============================================
doc.add_heading('7. Troubleshooting Guide', level=1)
doc.add_paragraph('If you encounter issues, refer to the table below for possible causes and solutions:')

issues = [
    ['Cannot access the application', 'Docker services are not running', 'Run "docker compose up -d" in the project directory and wait for all containers to start.'],
    ['Login fails with "Invalid email or password"', 'Incorrect credentials or inactive account', 'Verify your email and password. Contact your manager to ensure your account is active.'],
    ['Login fails with "Account is inactive"', 'Your account has been deactivated', 'Contact your supervisor or manager to reactivate your account.'],
    ['Data not syncing after coming online', 'Network or server connectivity issue', 'Check your internet connection. Verify the backend server is running. Click the sync status indicator to manually trigger a sync.'],
    ['Citizen registration fails to save', 'Offline storage full or browser issue', 'Clear browser cache and try again. Ensure your browser supports IndexedDB. Check available disk space.'],
    ['GPS location not captured', 'Location permissions denied in browser', 'Grant location permissions in your browser settings. Ensure GPS is enabled on your device.'],
    ['Photos not uploading', 'File size too large or server storage full', 'Reduce photo file size before uploading. Contact your administrator to check server storage.'],
    ['Dashboard shows no data', 'No records exist yet or sync pending', 'Ensure data has been entered and synced. Check the sync status indicator for pending items.'],
    ['Email notifications not received', 'Email server not configured', 'Verify the EMAIL_USER and EMAIL_PASS environment variables are correctly set in the .env file.'],
    ['Application is slow', 'Server resource constraints', 'Check server CPU and memory usage. Consider upgrading resources or restarting the Docker containers.'],
    ['Screen time not tracking', 'Screen time monitoring disabled or blocked', 'Ensure the screen time feature is enabled in the application settings. Check browser permissions.'],
    ['Cannot change password', 'Incorrect current password', 'Verify your current password before entering a new one. Passwords must meet minimum security requirements.'],
]
make_table(doc, ['Issue', 'Possible Cause', 'Solution'], issues)

# =============================================
# 8. FAQ
# =============================================
doc.add_heading('8. Frequently Asked Questions (FAQ)', level=1)

faqs = [
    ('Q: Can I use FieldSync on my mobile phone?',
     'A: Yes. FieldSync uses a responsive web design that adapts to mobile screens. You can also install it as a Progressive Web App (PWA) for a native app-like experience on your phone.'),
    ('Q: What happens to my data if I lose internet connection?',
     'A: All data entered while offline is saved locally in your browser using IndexedDB. When your connection is restored, the data is automatically synced to the server. No data is lost.'),
    ('Q: How do I know if my data has been synced?',
     'A: The sync status indicator in the header shows your current connectivity status. A badge counter displays the number of items pending sync. When all items are synced, the counter clears.'),
    ('Q: Can I use FieldSync in my local language?',
     'A: Yes. FieldSync supports English, Amharic (Amharic), Tigrinya (Tigrinya), and Oromo (Oromo). You can switch languages at any time from the language selector in the sidebar or header.'),
    ('Q: What should I do if I forget my password?',
     'A: Contact your supervisor or manager. They can assist with password reset or contact the system administrator for account recovery.'),
    ('Q: How is my GPS location used?',
     'A: GPS coordinates are captured during citizen registration and optionally during field reports. This helps track field coverage and ensures accurate location-based data. Location data is only visible to supervisors and managers.'),
    ('Q: Can I edit a report after submitting it?',
     'A: Once a report is synced to the server, it cannot be directly edited. If corrections are needed, create a supplementary report or contact your supervisor.'),
    ('Q: What is the verification system?',
     'A: The verification system periodically prompts you to confirm your identity. This ensures that the authorized user is actively using the system. Missing a verification alert notifies your supervisor.'),
    ('Q: How do I upload a profile photo?',
     'A: Navigate to your Profile page from the sidebar, click on the avatar area, and select a photo from your device. Supported formats include JPG and PNG.'),
    ('Q: Is my data secure?',
     'A: Yes. FieldSync uses secure password hashing (bcrypt), JWT-based authentication, encrypted data transmission, and role-based access control. Server data is backed up via Docker volumes.'),
]

for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph(a)

# =============================================
# 9. SUPPORT AND CONTACT
# =============================================
doc.add_heading('9. Support and Contact', level=1)
doc.add_paragraph('If you need further assistance, please contact our support team:')

doc.add_paragraph()
support_items = [
    ['Email', 'support@fieldsync.gov'],
    ['Phone', '+251-11-000-0000'],
    ['Website', 'https://fieldsync.gov/support'],
]
make_table(doc, ['Channel', 'Details'], support_items)

doc.add_paragraph()
doc.add_paragraph(
    'Live chat support is available at: https://fieldsync.gov/livechat'
)
doc.add_paragraph(
    'Support Hours: Monday to Friday, 8:00 AM - 5:00 PM (Local Time). '
    'For urgent issues outside business hours, please email support@fieldsync.gov with '
    'subject line prefixed with [URGENT].'
)

# =============================================
# 10. REVISION HISTORY
# =============================================
doc.add_heading('10. Revision History', level=1)

revisions = [
    ['1.0.0', datetime.date.today().strftime('%Y-%m-%d'), 'Initial release of the FieldSync User Manual.'],
]
make_table(doc, ['Version', 'Date', 'Changes'], revisions)

# =============================================
# SAVE
# =============================================
output_path = 'C:\\Users\\nebi\\Desktop\\mongoreact\\fieldsync\\FieldSync_User_Manual_v2.docx'
doc.save(output_path)
print(f'User Manual saved to: {output_path}')
