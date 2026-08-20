"""
FieldSync System Design Document Generator
SDD-only with attractive diagrams created via Pillow.
Same styling as existing documents (Calibri, navy headers, Light Grid tables).
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont
import datetime
import os
import math

# ─── COLOR PALETTE ───
NAVY = (30, 58, 95)
NAVY_HEX = '1E3A5F'
LIGHT_BLUE = (79, 195, 247)
DARK_BLUE = (13, 71, 161)
TEAL = (0, 150, 136)
GREEN = (76, 175, 80)
ORANGE = (255, 152, 0)
RED = (211, 47, 47)
GRAY = (158, 158, 158)
WHITE = (255, 255, 255)
BLACK = (33, 33, 33)
PURPLE = (123, 31, 162)
AMBER = (255, 193, 7)

OUT_DIR = r'C:\Users\nebi\Desktop\mongoreact\fieldsync'

# ─── DRAWING HELPERS ───
def try_font(size, bold=False):
    names = ['arialbd.ttf', 'arial.ttf', 'segoeui.ttf', 'seguisb.ttf']
    if bold:
        names = ['arialbd.ttf', 'seguisb.ttf'] + names
    for n in names:
        try:
            return ImageFont.truetype(n, size)
        except:
            pass
    return ImageFont.load_default()

def rounded_rect(draw, xy, radius, fill=None, outline=None, width=1):
    x0, y0, x1, y1 = xy
    r = radius
    if fill:
        draw.rectangle([x0+r, y0, x1-r, y1], fill=fill)
        draw.rectangle([x0, y0+r, x1, y1-r], fill=fill)
        draw.pieslice([x0, y0, x0+2*r, y0+2*r], 180, 270, fill=fill)
        draw.pieslice([x1-2*r, y0, x1, y0+2*r], 270, 360, fill=fill)
        draw.pieslice([x0, y1-2*r, x0+2*r, y1], 90, 180, fill=fill)
        draw.pieslice([x1-2*r, y1-2*r, x1, y1], 0, 90, fill=fill)
    if outline:
        draw.arc([x0, y0, x0+2*r, y0+2*r], 180, 270, fill=outline, width=width)
        draw.arc([x1-2*r, y0, x1, y0+2*r], 270, 360, fill=outline, width=width)
        draw.arc([x0, y1-2*r, x0+2*r, y1], 90, 180, fill=outline, width=width)
        draw.arc([x1-2*r, y1-2*r, x1, y1], 0, 90, fill=outline, width=width)
        draw.line([x0+r, y0, x1-r, y0], fill=outline, width=width)
        draw.line([x0+r, y1, x1-r, y1], fill=outline, width=width)
        draw.line([x0, y0+r, x0, y1-r], fill=outline, width=width)
        draw.line([x1, y0+r, x1, y1-r], fill=outline, width=width)

def draw_box(draw, x, y, w, h, text, fill_color, text_color=WHITE, font_size=11, radius=10, outline_color=None, subtext=None):
    rounded_rect(draw, (x, y, x+w, y+h), radius, fill=fill_color, outline=outline_color or fill_color, width=2)
    font = try_font(font_size, bold=True)
    small_font = try_font(font_size - 3)
    bbox = draw.textbbox((0, 0), text, font=font)
    tw = bbox[2] - bbox[0]
    th = bbox[3] - bbox[1]
    if subtext:
        sbbox = draw.textbbox((0, 0), subtext, font=small_font)
        sth = sbbox[3] - sbbox[1]
        total_h = th + sth + 4
        ty = y + (h - total_h) // 2
        draw.text((x + (w - tw) // 2, ty), text, fill=text_color, font=font)
        sbw = sbbox[2] - sbbox[0]
        draw.text((x + (w - sbw) // 2, ty + th + 4), subtext, fill=text_color, font=small_font)
    else:
        ty = y + (h - th) // 2
        draw.text((x + (w - tw) // 2, ty), text, fill=text_color, font=font)

def draw_arrow_down(draw, x1, y1, x2, y2, color, width=2, head_size=8):
    draw.line([(x1, y1), (x2, y2)], fill=color, width=width)
    angle = math.atan2(y2 - y1, x2 - x1)
    la = angle + math.pi * 0.8
    ra = angle - math.pi * 0.8
    draw.polygon([
        (x2, y2),
        (x2 + head_size * math.cos(la), y2 + head_size * math.sin(la)),
        (x2 + head_size * math.cos(ra), y2 + head_size * math.sin(ra)),
    ], fill=color)

def draw_right_arrow(draw, x1, y, x2, color, width=2, head_size=7):
    draw.line([(x1, y), (x2, y)], fill=color, width=width)
    draw.polygon([(x2, y), (x2 - head_size, y - head_size//2), (x2 - head_size, y + head_size//2)], fill=color)

def draw_dashed_line(draw, x1, y1, x2, y2, color, width=2, dash=8, gap=5):
    dx = x2 - x1
    dy = y2 - y1
    dist = math.sqrt(dx*dx + dy*dy)
    if dist == 0:
        return
    ux, uy = dx/dist, dy/dist
    pos = 0
    while pos < dist:
        sx = x1 + ux * pos
        sy = y1 + uy * pos
        end = min(pos + dash, dist)
        ex = x1 + ux * end
        ey = y1 + uy * end
        draw.line([(sx, sy), (ex, ey)], fill=color, width=width)
        pos = end + gap

def draw_label(draw, x, y, text, font_size=9, color=(100, 100, 100)):
    font = try_font(font_size)
    bbox = draw.textbbox((0, 0), text, font=font)
    tw = bbox[2] - bbox[0]
    draw.text((x - tw // 2, y), text, fill=color, font=font)

def draw_section_bg(draw, x, y, w, h, outline_color, fill_color):
    rounded_rect(draw, (x, y, x+w, y+h), 12, fill=fill_color, outline=outline_color, width=2)


# ═══════════════════════════════════════════════════════
# DIAGRAM 1: SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════
def create_architecture_diagram():
    W, H = 1400, 1000
    img = Image.new('RGB', (W, H), WHITE)
    draw = ImageDraw.Draw(img)
    rounded_rect(draw, (8, 8, W-8, H-8), 15, fill=None, outline=(180, 180, 180), width=1)

    # Title
    draw.text((W//2 - 180, 18), 'FieldSync System Architecture', fill=NAVY, font=try_font(22, bold=True))
    draw.text((W//2 - 120, 48), 'Three-Tier Client-Server Architecture', fill=(120, 120, 120), font=try_font(12))

    # ── CLIENT TIER ──
    draw_section_bg(draw, 30, 75, W-60, 250, LIGHT_BLUE, (232, 244, 253))
    draw_label(draw, W//2, 80, 'CLIENT BROWSER (PWA)', 13, NAVY)

    bw, bh = 200, 80
    cx = 100
    draw_box(draw, cx, 115, bw, bh, 'React SPA', LIGHT_BLUE, WHITE, 13, subtext='UI / Routing / i18next')
    draw_box(draw, cx + 260, 115, bw, bh, 'Dexie.js', DARK_BLUE, WHITE, 13, subtext='IndexedDB Storage')
    draw_box(draw, cx + 520, 115, bw, bh, 'Sync Service', (0, 121, 107), WHITE, 13, subtext='Offline Queue + Auto Sync')

    draw_box(draw, cx + 100, 215, 180, 70, 'PWA / Service Worker', PURPLE, WHITE, 10, radius=8)
    draw_box(draw, cx + 330, 215, 180, 70, 'React Hook Form + Zod', RED, WHITE, 10, radius=8)
    draw_box(draw, cx + 560, 215, 180, 70, 'Recharts / i18next', ORANGE, WHITE, 10, radius=8)

    # Docker Compose label
    draw_box(draw, W-180, 85, 155, 40, 'Docker Compose', NAVY, WHITE, 10, radius=6)
    draw_box(draw, W-180, 130, 155, 35, 'db | backend | frontend', GRAY, WHITE, 8, radius=6)

    # Arrows client -> nginx
    for ax in [300, 560, 820]:
        draw_dashed_line(draw, ax, 295, ax, 340, GRAY, width=2)
    draw_label(draw, 560, 300, 'HTTP/REST (JSON)', 10, NAVY)
    draw_arrow_down(draw, 560, 330, 560, 345, GRAY, 2, 7)

    # ── NGINX TIER ──
    draw_section_bg(draw, 30, 345, W-60, 80, GRAY, (238, 238, 238))
    draw_box(draw, 150, 360, W-300, 50, 'NGINX REVERSE PROXY    (Port 80 \u2192 Frontend Static / API Proxy)', (117, 117, 117), WHITE, 13)

    # Arrow nginx -> backend
    draw_dashed_line(draw, 560, 425, 560, 470, GRAY, width=2)
    draw_label(draw, 560, 430, 'TCP (HTTP)', 10, NAVY)
    draw_arrow_down(draw, 560, 460, 560, 475, GRAY, 2, 7)

    # ── BACKEND TIER ──
    draw_section_bg(draw, 30, 475, W-60, 310, NAVY, (232, 240, 251))
    draw_label(draw, W//2, 480, 'NODE.JS / EXPRESS API SERVER (TypeScript)', 13, NAVY)

    modules_r1 = [
        ('Auth Module', RED), ('Citizen Module', (21, 101, 192)),
        ('Report Module', TEAL), ('Task / Permission', ORANGE),
    ]
    modules_r2 = [
        ('User Module', PURPLE), ('Audit Module', GRAY),
        ('Sync Module', (0, 150, 136)), ('File Upload (Multer)', PURPLE),
    ]

    bw2, bh2 = 180, 55
    start_x = 80
    gap = 230
    for i, (name, color) in enumerate(modules_r1):
        draw_box(draw, start_x + i * gap, 510, bw2, bh2, name, color, WHITE, 11, radius=8)
    for i, (name, color) in enumerate(modules_r2):
        draw_box(draw, start_x + i * gap, 580, bw2, bh2, name, color, WHITE, 11, radius=8)

    draw_box(draw, 100, 655, W-200, 45, 'Nodemailer (Email Notifications \u2014 SMTP/TLS)', (198, 40, 40), WHITE, 12, radius=8)

    draw_box(draw, 100, 715, 200, 35, 'bcrypt + JWT Auth', NAVY, WHITE, 10, radius=6)
    draw_box(draw, 350, 715, 200, 35, 'Zod Validation', TEAL, WHITE, 10, radius=6)
    draw_box(draw, 600, 715, 200, 35, 'CORS Middleware', GRAY, WHITE, 10, radius=6)
    draw_box(draw, 850, 715, 200, 35, 'RBAC Middleware', RED, WHITE, 10, radius=6)

    # Arrow backend -> db
    draw_dashed_line(draw, 560, 785, 560, 830, GRAY, width=2)
    draw_label(draw, 560, 790, 'TCP (pg driver)', 10, NAVY)
    draw_arrow_down(draw, 560, 820, 560, 835, GRAY, 2, 7)

    # ── DATABASE TIER ──
    draw_section_bg(draw, 30, 835, W-60, 145, TEAL, (224, 242, 241))
    draw_label(draw, W//2, 840, 'PostgreSQL 16 (Alpine) \u2014 Docker Volume', 13, (0, 100, 80))

    tables = ['users', 'citizens', 'reports', 'tasks', 'permissions',
              'alerts', 'audit', 'screen_time', 'verification', 'locations', 'supervisor_reports']
    tw, th = 110, 32
    for i, tbl in enumerate(tables):
        row_idx = i // 5
        col_idx = i % 5
        tx = 60 + col_idx * 118
        ty = 865 + row_idx * 45
        rounded_rect(draw, (tx, ty, tx+tw, ty+th), 5, fill=TEAL, outline=(0, 100, 80), width=1)
        font = try_font(9, bold=True)
        bbox = draw.textbbox((0, 0), tbl, font=font)
        ttw = bbox[2] - bbox[0]
        draw.text((tx + (tw - ttw) // 2, ty + 8), tbl, fill=WHITE, font=font)

    path = os.path.join(OUT_DIR, 'diagram_architecture.png')
    img.save(path, quality=95)
    return path


# ═══════════════════════════════════════════════════════
# DIAGRAM 2: DATA FLOW - CITIZEN REGISTRATION
# ═══════════════════════════════════════════════════════
def create_citizen_flow_diagram():
    W, H = 1200, 750
    img = Image.new('RGB', (W, H), WHITE)
    draw = ImageDraw.Draw(img)
    rounded_rect(draw, (8, 8, W-8, H-8), 15, fill=None, outline=(180, 180, 180), width=1)
    draw.text((W//2 - 200, 18), 'Citizen Registration Data Flow', fill=NAVY, font=try_font(20, bold=True))

    bw, bh = 170, 65

    # ── ONLINE PATH (left column) ──
    draw_box(draw, 30, 80, 170, 60, 'Field Officer', NAVY, WHITE, 14)
    draw_arrow_down(draw, 115, 140, 115, 170, NAVY, 2, 7)

    draw_box(draw, 30, 175, 170, 60, 'React Form', LIGHT_BLUE, WHITE, 13, subtext='Zod + React Hook Form')
    draw_arrow_down(draw, 115, 235, 115, 265, LIGHT_BLUE, 2, 7)

    draw_box(draw, 30, 270, 170, 60, 'GPS Capture', TEAL, WHITE, 12)
    draw_arrow_down(draw, 115, 330, 115, 360, TEAL, 2, 7)

    draw_box(draw, 30, 365, 170, 55, 'Validation OK?', ORANGE, WHITE, 12)

    # YES branch
    draw_label(draw, 265, 380, 'YES (Online)', 11, GREEN)
    draw.line([(200, 392), (250, 392)], fill=GREEN, width=2)
    draw_arrow_down(draw, 250, 392, 340, 392, GREEN, 2, 7)
    draw_arrow_down(draw, 340, 392, 340, 415, GREEN, 2, 7)

    draw_box(draw, 270, 415, 170, 55, 'POST /api/citizens', (21, 101, 192), WHITE, 12)
    draw_arrow_down(draw, 355, 470, 355, 500, (21, 101, 192), 2, 7)

    draw_box(draw, 270, 505, 170, 55, 'Express Backend', NAVY, WHITE, 12, subtext='Validate + Process')
    draw_arrow_down(draw, 355, 560, 355, 590, NAVY, 2, 7)

    draw_box(draw, 270, 595, 170, 55, 'PostgreSQL', TEAL, WHITE, 13, subtext='citizens table')

    # NO branch (Offline)
    draw_label(draw, 115, 430, 'NO (Offline)', 11, ORANGE)
    draw_dashed_line(draw, 115, 420, 115, 465, ORANGE, 2, 6, 4)
    draw_arrow_down(draw, 115, 455, 115, 470, ORANGE, 2, 7)

    draw_box(draw, 20, 475, 195, 60, 'Dexie.js / IndexedDB', PURPLE, WHITE, 12, subtext='Local Storage')
    draw_arrow_down(draw, 117, 535, 117, 565, PURPLE, 2, 7)

    draw_box(draw, 20, 570, 195, 55, 'Sync Queue (Pending)', ORANGE, WHITE, 12)
    draw_arrow_down(draw, 117, 625, 117, 655, ORANGE, 2, 7)

    draw_box(draw, 20, 660, 195, 55, 'NetworkStatus Hook', TEAL, WHITE, 11, subtext='Detects Reconnection')

    # Reconnect arrow
    draw_label(draw, 500, 540, 'Reconnects & Syncs', 11, PURPLE)
    draw.line([(215, 687), (260, 687)], fill=PURPLE, width=2)
    draw.line([(260, 687), (260, 540)], fill=PURPLE, width=2)
    draw_arrow_down(draw, 260, 540, 270, 540, PURPLE, 2, 6)

    # ── RIGHT SIDE: Audit + Response ──
    draw_box(draw, 520, 505, 170, 55, 'Audit Middleware', GRAY, WHITE, 12)
    draw_arrow_down(draw, 605, 560, 605, 590, GRAY, 2, 7)
    draw_box(draw, 520, 595, 170, 55, 'audit table', GRAY, WHITE, 12)
    draw.line([(440, 532), (520, 532)], fill=GRAY, width=2)

    draw_box(draw, 520, 680, 170, 50, 'Manager Reviews', NAVY, WHITE, 11, subtext='Compliance Audit')
    draw_dashed_line(draw, 605, 650, 605, 680, GRAY, 2, 6, 4)

    draw_box(draw, 750, 505, 190, 55, 'Response (JSON)', TEAL, WHITE, 12)
    draw_arrow_down(draw, 845, 560, 845, 590, TEAL, 2, 7)
    draw_box(draw, 750, 595, 190, 55, 'React Render', TEAL, WHITE, 12, subtext='Update UI')
    draw.line([(440, 532), (750, 532)], fill=TEAL, width=2)

    # Legend
    lx, ly = 750, 100
    rounded_rect(draw, (lx, ly, lx+380, ly+250), 10, fill=(248, 248, 248), outline=(200, 200, 200), width=1)
    draw.text((lx+10, ly+8), 'Legend', fill=NAVY, font=try_font(13, bold=True))
    for i, (c, t) in enumerate([
        ((21, 101, 192), 'Online Path'), (ORANGE, 'Offline Path'),
        (PURPLE, 'Sync/Reconnect Path'), (GREEN, 'Data Flow'),
        (GRAY, 'Audit Path'), (TEAL, 'Response Path'),
    ]):
        yy = ly + 35 + i * 32
        rounded_rect(draw, (lx+15, yy, lx+45, yy+18), 4, fill=c)
        draw.text((lx+55, yy), t, fill=BLACK, font=try_font(11))

    path = os.path.join(OUT_DIR, 'diagram_citizen_flow.png')
    img.save(path, quality=95)
    return path


# ═══════════════════════════════════════════════════════
# DIAGRAM 3: AUTHENTICATION FLOW
# ═══════════════════════════════════════════════════════
def create_auth_flow_diagram():
    W, H = 1100, 620
    img = Image.new('RGB', (W, H), WHITE)
    draw = ImageDraw.Draw(img)
    rounded_rect(draw, (8, 8, W-8, H-8), 15, fill=None, outline=(180, 180, 180), width=1)
    draw.text((W//2 - 180, 18), 'Authentication Data Flow', fill=NAVY, font=try_font(20, bold=True))

    bw, bh = 150, 55
    cols = [60, 250, 470, 700, 900]

    # Row 1
    draw_box(draw, 60, 80, 150, 55, 'User', NAVY, WHITE, 14)
    draw_box(draw, 250, 80, 170, 55, 'Login Page', LIGHT_BLUE, WHITE, 13, subtext='Email + Password')
    draw_box(draw, 470, 80, 190, 55, 'POST /api/auth/login', (21, 101, 192), WHITE, 11)
    draw_box(draw, 700, 80, 170, 55, 'Express Backend', NAVY, WHITE, 13)
    draw_box(draw, 900, 80, 150, 55, 'PostgreSQL', TEAL, WHITE, 13)

    for x1, x2 in [(210, 250), (420, 470), (660, 700), (870, 900)]:
        draw.line([(x1, 107), (x2, 107)], fill=NAVY, width=2)
        draw.polygon([(x2, 107), (x2-7, 103), (x2-7, 111)], fill=NAVY)

    # Step boxes
    draw_label(draw, 500, 150, 'Step-by-step process', 12, NAVY)
    steps = [
        ('1. Enter email\n    + password', cols[0], RED),
        ('2. Send POST\n    request', cols[1], (21, 101, 192)),
        ('3. Retrieve user\n    by email', cols[2], NAVY),
        ('4. Query users\n    table', cols[3], TEAL),
    ]
    for i, (text, x, c) in enumerate(steps):
        draw_box(draw, x, 180, 160, 60, text, c, WHITE, 10)
        if i < len(steps) - 1:
            draw.line([(x+160, 210), (steps[i+1][1], 210)], fill=GRAY, width=2)
            draw.polygon([(steps[i+1][1], 210), (steps[i+1][1]-7, 206), (steps[i+1][1]-7, 214)], fill=GRAY)

    # Row 3
    draw_box(draw, 100, 290, 200, 60, '5. Compare bcrypt\n    hash', RED, WHITE, 11)
    draw_box(draw, 370, 290, 200, 60, '6. Generate\n    JWT Token', ORANGE, WHITE, 11)
    draw_box(draw, 640, 290, 220, 60, '7. Return user + token\n    (JSON response)', TEAL, WHITE, 11)
    draw.line([(300, 320), (370, 320)], fill=GRAY, width=2)
    draw.polygon([(370, 320), (363, 316), (363, 324)], fill=GRAY)
    draw.line([(570, 320), (640, 320)], fill=GRAY, width=2)
    draw.polygon([(640, 320), (633, 316), (633, 324)], fill=GRAY)

    # Row 4
    draw_box(draw, 200, 390, 250, 60, '8. Store token in\n    React Context', PURPLE, WHITE, 11)
    draw_box(draw, 550, 390, 300, 60, '9. All subsequent requests include\n    Authorization: Bearer <token>', GRAY, WHITE, 11)
    draw_arrow_down(draw, 750, 350, 750, 390, TEAL, 2, 7)
    draw_arrow_down(draw, 325, 350, 325, 390, PURPLE, 2, 7)
    draw.line([(450, 420), (550, 420)], fill=GRAY, width=2)
    draw.polygon([(550, 420), (543, 416), (543, 424)], fill=GRAY)

    # Security badges
    draw_box(draw, 50, 490, 220, 45, 'bcrypt (10 salt rounds)', RED, WHITE, 10, radius=6)
    draw_box(draw, 300, 490, 220, 45, 'JWT (HS256 signing)', ORANGE, WHITE, 10, radius=6)
    draw_box(draw, 550, 490, 220, 45, 'CORS Protection', GRAY, WHITE, 10, radius=6)
    draw_box(draw, 800, 490, 220, 45, 'RBAC Middleware', TEAL, WHITE, 10, radius=6)

    # Note
    rounded_rect(draw, (50, 555, W-50, 595), 8, fill=(255, 248, 225), outline=AMBER, width=1)
    draw.text((70, 565), 'Note: Tokens have configurable expiration. Refresh required for extended sessions. '
              'Environment variable JWT_SECRET must be a strong, unique string.', fill=(120, 80, 0), font=try_font(10))

    path = os.path.join(OUT_DIR, 'diagram_auth_flow.png')
    img.save(path, quality=95)
    return path


# ═══════════════════════════════════════════════════════
# DIAGRAM 4: SYNC / OFFLINE FLOW
# ═══════════════════════════════════════════════════════
def create_sync_flow_diagram():
    W, H = 1100, 600
    img = Image.new('RGB', (W, H), WHITE)
    draw = ImageDraw.Draw(img)
    rounded_rect(draw, (8, 8, W-8, H-8), 15, fill=None, outline=(180, 180, 180), width=1)
    draw.text((W//2 - 180, 18), 'Offline Sync Data Flow', fill=NAVY, font=try_font(20, bold=True))

    bw, bh = 200, 55

    # LEFT: Online
    draw_section_bg(draw, 20, 60, 500, 500, TEAL, (224, 242, 241))
    draw_label(draw, 270, 65, 'ONLINE PATH', 13, (0, 100, 80))
    online = [
        (180, 95, 'User Action', NAVY, 'Create/Edit Data'),
        (180, 185, 'Axios HTTP Request', (21, 101, 192), 'POST/PUT/DELETE'),
        (180, 275, 'Express API', NAVY, 'Controller + Model'),
        (180, 365, 'PostgreSQL', TEAL, 'Persistent Storage'),
        (180, 455, 'Response + UI Update', TEAL, 'Success/Error Toast'),
    ]
    for (x, y, text, color, sub) in online:
        draw_box(draw, x, y, bw, bh, text, color, WHITE, 12, subtext=sub)
    for i in range(len(online) - 1):
        draw_arrow_down(draw, 180+bw//2, online[i][1]+bh+2, 180+bw//2, online[i+1][1]-2, GREEN, 2, 6)

    # RIGHT: Offline
    draw_section_bg(draw, 560, 60, 500, 500, ORANGE, (255, 243, 224))
    draw_label(draw, 810, 65, 'OFFLINE PATH', 13, (200, 100, 0))
    offline = [
        (700, 95, 'User Action', NAVY, 'Create/Edit Data'),
        (700, 185, 'SyncService Intercept', ORANGE, 'Check Network Status'),
        (700, 275, 'Dexie.js / IndexedDB', PURPLE, 'Local Storage'),
        (700, 365, 'Pending Sync Queue', ORANGE, 'Status: "pending"'),
        (700, 455, 'NetworkStatus Hook', TEAL, 'Detects Reconnection'),
    ]
    for (x, y, text, color, sub) in offline:
        draw_box(draw, x, y, bw, bh, text, color, WHITE, 12, subtext=sub)
    for i in range(len(offline) - 1):
        draw_arrow_down(draw, 700+bw//2, offline[i][1]+bh+2, 700+bw//2, offline[i+1][1]-2, ORANGE, 2, 6)

    # Reconnect arrow
    draw_label(draw, 500, 480, 'Reconnect', 11, PURPLE)
    cx_off = 700 + bw // 2
    draw.line([(cx_off, 455+bh), (cx_off, 535)], fill=PURPLE, width=2)
    draw.line([(cx_off, 535), (500, 535)], fill=PURPLE, width=2)
    draw.line([(500, 535), (500, 400)], fill=PURPLE, width=2)
    draw_arrow_down(draw, 500, 400, 500, 395, PURPLE, 2, 6)
    draw_label(draw, 520, 520, 'Replay All Pending Ops', 10, PURPLE)

    draw_box(draw, 20, 570, 280, 25, 'Sync Status Indicator: Online / Offline / Pending count', GRAY, WHITE, 9, radius=5)

    path = os.path.join(OUT_DIR, 'diagram_sync_flow.png')
    img.save(path, quality=95)
    return path


# ═══════════════════════════════════════════════════════
# DIAGRAM 5: ER / ENTITY RELATIONSHIP
# ═══════════════════════════════════════════════════════
def create_er_diagram():
    W, H = 1300, 750
    img = Image.new('RGB', (W, H), WHITE)
    draw = ImageDraw.Draw(img)
    rounded_rect(draw, (8, 8, W-8, H-8), 15, fill=None, outline=(180, 180, 180), width=1)
    draw.text((W//2 - 150, 18), 'Data Entity Relationships', fill=NAVY, font=try_font(20, bold=True))

    entities = [
        {'name': 'users', 'x': 30, 'y': 70, 'color': (21, 101, 192),
         'fields': ['id (PK)', 'email', 'password_hash', 'role', 'team_id', 'language', 'status']},
        {'name': 'citizens', 'x': 290, 'y': 70, 'color': TEAL,
         'fields': ['id (PK)', 'first_name', 'last_name', 'dob', 'gender', 'region', 'gps_lat', 'gps_long']},
        {'name': 'reports', 'x': 550, 'y': 70, 'color': PURPLE,
         'fields': ['id (PK)', 'title', 'description', 'officer_id (FK)', 'sync_status', 'photo_url']},
        {'name': 'tasks', 'x': 810, 'y': 70, 'color': RED,
         'fields': ['id (PK)', 'title', 'description', 'assignee_id (FK)', 'priority', 'due_date']},
        {'name': 'permissions', 'x': 30, 'y': 350, 'color': ORANGE,
         'fields': ['id (PK)', 'type', 'reason', 'requester_id (FK)', 'approver_id (FK)', 'status']},
        {'name': 'alerts', 'x': 290, 'y': 350, 'color': PURPLE,
         'fields': ['id (PK)', 'sender_id (FK)', 'recipient_id (FK)', 'subject', 'message', 'is_read']},
        {'name': 'audit', 'x': 550, 'y': 350, 'color': GRAY,
         'fields': ['id (PK)', 'user_id (FK)', 'action', 'entity_type', 'entity_id', 'metadata']},
        {'name': 'screen_time', 'x': 810, 'y': 350, 'color': TEAL,
         'fields': ['id (PK)', 'user_id (FK)', 'start_time', 'end_time', 'duration', 'device']},
        {'name': 'verification', 'x': 30, 'y': 580, 'color': RED,
         'fields': ['id (PK)', 'user_id (FK)', 'prompt', 'response', 'timestamp', 'status']},
        {'name': 'locations', 'x': 290, 'y': 580, 'color': TEAL,
         'fields': ['id (PK)', 'name', 'type', 'parent_id (FK)', 'region', 'district']},
        {'name': 'supervisor_reports', 'x': 550, 'y': 580, 'color': PURPLE,
         'fields': ['id (PK)', 'supervisor_id (FK)', 'period', 'metrics', 'summary', 'created_at']},
    ]

    ew, eh = 230, 140
    header_h = 28
    for ent in entities:
        x, y, c = ent['x'], ent['y'], ent['color']
        rounded_rect(draw, (x, y, x+ew, y+header_h), 8, fill=c, outline=c)
        rounded_rect(draw, (x, y+header_h, x+ew, y+eh), 8, fill=WHITE, outline=c, width=2)
        draw.rectangle([x+1, y+header_h, x+ew-1, y+header_h+1], fill=c)
        font = try_font(12, bold=True)
        bbox = draw.textbbox((0, 0), ent['name'], font=font)
        tw = bbox[2] - bbox[0]
        draw.text((x + (ew - tw) // 2, y + 6), ent['name'], fill=WHITE, font=font)
        field_font = try_font(9)
        for i, f in enumerate(ent['fields']):
            fy = y + header_h + 8 + i * 16
            color = BLACK if 'FK' not in f else (150, 100, 50)
            draw.text((x + 10, fy), f, fill=color, font=field_font)

    # Relationship lines
    draw.line([(30+ew, 70+40), (550, 70+40)], fill=(21, 101, 192), width=1)
    draw_label(draw, 400, 55, '1:N (officer_id)', 8, (21, 101, 192))
    draw.line([(30+ew, 70+60), (810, 70+60)], fill=RED, width=1)
    draw_label(draw, 600, 55, '1:N (assignee_id)', 8, RED)
    draw.line([(30+ew//2, 70+eh), (30+ew//2, 350)], fill=ORANGE, width=1)
    draw_label(draw, 15, 280, '1:N', 8, ORANGE)
    draw.line([(30+ew, 70+80), (290, 350+30)], fill=PURPLE, width=1)
    draw_label(draw, 200, 230, '1:N', 8, PURPLE)
    draw.line([(30+ew, 70+100), (550, 350+30)], fill=GRAY, width=1)
    draw_label(draw, 400, 250, '1:N', 8, GRAY)
    draw.line([(30+ew, 70+120), (810, 350+30)], fill=TEAL, width=1)
    draw.line([(30+ew//2, 350), (30+ew//2, 580)], fill=RED, width=1)
    draw_label(draw, 80, 480, '1:N', 8, RED)

    # Legend
    rounded_rect(draw, (1050, 70, 1280, 200), 8, fill=(248, 248, 248), outline=(200, 200, 200))
    draw.text((1060, 80), 'FK = Foreign Key', fill=(150, 100, 50), font=try_font(10, bold=True))
    draw.text((1060, 100), 'PK = Primary Key', fill=(21, 101, 192), font=try_font(10, bold=True))
    draw.text((1060, 120), '1:N = One-to-Many', fill=GRAY, font=try_font(10))
    draw.text((1060, 140), 'All entities link back', fill=GRAY, font=try_font(10))
    draw.text((1060, 155), 'to the users table', fill=GRAY, font=try_font(10))

    path = os.path.join(OUT_DIR, 'diagram_er.png')
    img.save(path, quality=95)
    return path


# ═══════════════════════════════════════════════════════
# WORD DOCUMENT - SDD ONLY
# ═══════════════════════════════════════════════════════
def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells_data, bold=False, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        if bold or header:
            run.bold = True
        if header:
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_shading(cell, NAVY_HEX)
    return row

def make_table(doc, headers, rows):
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, headers, header=True)
    for row_data in rows:
        add_table_row(table, row_data)
    return table

def add_bullet_list(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def generate_sdd():
    doc = Document()

    # ── Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = RGBColor(*NAVY)
        hs.font.name = 'Calibri'

    # ═══════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('FieldSync')
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(*NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('System Design Document')
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(*LIGHT_BLUE)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Architecture, Components, Interfaces, and Data Flows')
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Version 1.0.0\n{datetime.date.today().strftime("%B %Y")}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════════════════
    doc.add_heading('Table of Contents', level=1)
    for item in [
        '1. Overview', '2. System Objectives', '3. Scope',
        '4. Assumptions and Constraints', '5. Stakeholders',
        '6. System Architecture', '7. Data Design', '8. Interfaces',
        '9. Security Considerations', '10. Performance Requirements',
        '11. Glossary', '12. Appendices',
    ]:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(4)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # 1. OVERVIEW
    # ═══════════════════════════════════════════════════════
    doc.add_heading('1. Overview', level=1)
    make_table(doc, ['Field', 'Details'], [
        ['System Name', 'FieldSync'],
        ['Prepared By', 'FieldSync Development Team'],
        ['Date', datetime.date.today().strftime('%Y-%m-%d')],
        ['Version', '1.0.0'],
        ['Description', 'A comprehensive field workforce management and citizen registration platform designed to streamline government field operations including National ID registration, reporting, task management, team oversight, and analytics.'],
    ])

    # ═══════════════════════════════════════════════════════
    # 2. SYSTEM OBJECTIVES
    # ═══════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_heading('2. System Objectives', level=1)
    add_bullet_list(doc, [
        'Enable field officers to register citizens for National ID programs accurately with GPS-verified location data and document capture.',
        'Provide offline-first architecture that ensures uninterrupted field operations regardless of internet connectivity, with automatic data synchronization upon reconnection.',
        'Establish a hierarchical role-based system (Field Officer, Supervisor, Manager) with granular access control for secure and efficient operations.',
        'Facilitate real-time task assignment, tracking, and completion workflows between supervisors and field officers.',
        'Deliver comprehensive analytics, audit trails, and reporting capabilities for managerial decision-making and regulatory compliance.',
        'Support multi-language accessibility (English, Amharic, Tigrinya, Oromo) to serve a diverse workforce.',
        'Ensure data integrity, security, and availability through encrypted storage, authentication, and automated backup mechanisms.',
        'Reduce paper-based processes and manual data entry by digitizing citizen registration and field reporting workflows.',
    ])

    # ═══════════════════════════════════════════════════════
    # 3. SCOPE
    # ═══════════════════════════════════════════════════════
    doc.add_heading('3. Scope', level=1)
    doc.add_heading('3.1 In Scope', level=2)
    add_bullet_list(doc, [
        'Citizen registration for National ID with biographic data, document types (National ID, Birth Certificate, Passport), GPS capture, and biometric collection status.',
        'User authentication and role-based authorization (Manager, Supervisor, Field Officer) with JWT tokens and bcrypt password hashing.',
        'Field reporting system with offline creation, photo/document attachment support, and automatic server synchronization.',
        'Task management workflow including creation, assignment, prioritization, due dates, and completion tracking.',
        'Team management for supervisors to monitor and manage field officer activities.',
        'Permission request and approval workflow for leaves, travel, and equipment requests.',
        'Supervisor report generation summarizing team performance and activities.',
        'Screen time tracking and identity verification for field officer monitoring.',
        'Alert and messaging system for internal notifications and team communication.',
        'Manager analytics dashboard with trend charts, registration statistics, and performance metrics.',
        'Audit logging of all system actions for accountability and compliance.',
        'Progressive Web App (PWA) support for mobile installation and offline access.',
        'Multi-language support with English, Amharic, Tigrinya, and Oromo translations.',
        'Dark and light theme toggle for user interface customization.',
        'Docker-based deployment with automated database migration and service orchestration.',
    ])
    doc.add_heading('3.2 Out of Scope', level=2)
    add_bullet_list(doc, [
        'Biometric data capture and storage (only collection status is tracked; actual biometric processing is handled by external systems).',
        'Payment processing or financial transaction management.',
        'Integration with external national ID issuance systems (data is prepared for export only).',
        'Mobile native applications (iOS/Android) \u2014 the system is web-based with PWA support.',
        'Video conferencing or real-time video streaming capabilities.',
    ])

    # ═══════════════════════════════════════════════════════
    # 4. ASSUMPTIONS AND CONSTRAINTS
    # ═══════════════════════════════════════════════════════
    doc.add_heading('4. Assumptions and Constraints', level=1)
    doc.add_heading('4.1 Assumptions', level=2)
    add_bullet_list(doc, [
        'Docker and Docker Compose are available and properly installed on the deployment server.',
        'The target deployment environment has stable power supply and reasonable network connectivity for initial setup.',
        'Field officers have access to modern web browsers (Chrome 90+, Firefox 88+, Edge 90+, or Safari 14+) on their devices.',
        'GPS and camera hardware is available on field officer devices for location capture and photo uploads.',
        'Administrators have basic command-line proficiency for Docker-based deployment and configuration.',
        'The PostgreSQL database server will be managed within Docker containers and does not require an external managed database service.',
        'Email services (SMTP) are available for notification delivery via Nodemailer.',
    ])
    doc.add_heading('4.2 Constraints', level=2)
    add_bullet_list(doc, [
        'The system requires Docker and Docker Compose for deployment; no native bare-metal installation path is provided.',
        "Offline data is stored in the browser's IndexedDB, which has device-dependent storage limits (typically 50 MB to several GB depending on browser and OS).",
        'All API communication uses HTTP/HTTPS; no WebSocket or real-time push notification infrastructure is currently in place.',
        'The frontend is served via Nginx in production; custom server-side rendering is not supported.',
        'File uploads (photos, documents) are limited to server storage capacity, which is backed by Docker volumes.',
        'Multi-language support covers the four specified languages; adding new languages requires manual translation updates.',
        'The system does not support multi-tenancy; each deployment serves a single organizational unit.',
    ])

    # ═══════════════════════════════════════════════════════
    # 5. STAKEHOLDERS
    # ═══════════════════════════════════════════════════════
    doc.add_heading('5. Stakeholders', level=1)
    make_table(doc, ['Name / Role', 'Responsibility', 'Contact'], [
        ['Project Sponsor', 'Program Director', 'sponsor@fieldsync.gov'],
        ['Project Manager', 'Delivery Lead', 'pm@fieldsync.gov'],
        ['Lead Developer', 'Backend & Architecture', 'lead-dev@fieldsync.gov'],
        ['Frontend Developer', 'UI/UX Implementation', 'frontend-dev@fieldsync.gov'],
        ['Backend Developer', 'API & Database', 'backend-dev@fieldsync.gov'],
        ['QA Engineer', 'Testing & Quality Assurance', 'qa@fieldsync.gov'],
        ['DevOps Engineer', 'Deployment & Infrastructure', 'devops@fieldsync.gov'],
        ['Field Officers', 'End Users (Registration & Reporting)', 'field-team@fieldsync.gov'],
        ['Supervisors', 'End Users (Team Management)', 'supervisors@fieldsync.gov'],
        ['Managers', 'End Users (Analytics & Administration)', 'managers@fieldsync.gov'],
        ['Database Administrator', 'Database Management & Backup', 'dba@fieldsync.gov'],
        ['Security Officer', 'Security & Compliance Review', 'security@fieldsync.gov'],
    ])

    # ═══════════════════════════════════════════════════════
    # 6. SYSTEM ARCHITECTURE
    # ═══════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_heading('6. System Architecture', level=1)
    doc.add_paragraph(
        'Overview: FieldSync follows a three-tier client-server architecture deployed via Docker Compose. '
        'The system comprises a React-based single-page application (SPA) frontend served by Nginx, '
        'a Node.js/Express REST API backend written in TypeScript, and a PostgreSQL 16 relational database. '
        'The frontend communicates with the backend exclusively through RESTful HTTP APIs. '
        'Offline-first capability is achieved via Dexie.js (IndexedDB) on the client side with an '
        'automatic synchronization layer that resolves conflicts upon reconnection.'
    )

    doc.add_heading('6.1 Components', level=2)
    make_table(doc, ['Component', 'Description'], [
        ['Frontend (React + Vite)',
         'A responsive Single Page Application built with React 18, Vite 5, and Tailwind CSS 3.4. '
         'Served by Nginx in production. Implements offline-first data storage using Dexie.js (IndexedDB wrapper), '
         'role-based UI rendering, multi-language support via i18next, form validation with React Hook Form + Zod, '
         'charts with Recharts, and PWA capabilities via vite-plugin-pwa.'],
        ['Backend API (Express + TypeScript)',
         'A RESTful API server built with Express 5 and TypeScript, following MVC architecture. '
         'Handles authentication (bcrypt + JWT), request validation, business logic, file uploads (Multer), '
         'and email notifications (Nodemailer). Runs on Node.js and exposes endpoints for auth, users, citizens, '
         'reports, tasks, permissions, alerts, audits, screen time, verification, locations, '
         'and supervisor reports.'],
        ['Database (PostgreSQL 16)',
         'Relational database using PostgreSQL 16 Alpine Docker image. Stores all persistent data including '
         'user accounts, citizen registrations, reports, tasks, permissions, alerts, audit logs, '
         'screen time data, and verification records. Data is persisted via Docker named volumes.'],
        ['Nginx (Reverse Proxy)',
         'Lightweight web server used to serve the frontend static build and proxy API requests to the backend. '
         'Configured as part of the frontend Docker container.'],
        ['Docker Compose',
         'Container orchestration layer that defines and manages three services (db, backend, frontend) with '
         'health checks, dependency ordering, volume mounts, and network configuration.'],
        ['Sync Service (Client-side)',
         'A client-side synchronization module (SyncService.js) that queues data operations in IndexedDB when '
         'offline and replays them against the backend API when connectivity is restored. Handles conflict resolution '
         'and provides sync status indicators.'],
        ['Offline Storage (Dexie.js / IndexedDB)',
         "Client-side database layer using Dexie.js as a wrapper around the browser's IndexedDB API. Stores "
         'citizen registrations, reports, and pending sync operations for offline access.'],
    ])

    doc.add_paragraph()
    doc.add_heading('6.2 System Architecture Diagram', level=2)
    doc.add_paragraph(
        'The system architecture follows a layered pattern with clear separation of concerns. '
        'The diagram below illustrates the three-tier architecture with all major components and their interactions:'
    )

    arch_path = os.path.join(OUT_DIR, 'diagram_architecture.png')
    if os.path.exists(arch_path):
        doc.add_picture(arch_path, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        'Figure 1: FieldSync System Architecture \u2014 Three-tier client-server architecture showing '
        'the Client Browser (PWA), Nginx Reverse Proxy, Node.js/Express API Server, and PostgreSQL 16 layers.'
    ).runs[0].font.size = Pt(9)

    # ═══════════════════════════════════════════════════════
    # 7. DATA DESIGN
    # ═══════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading('7. Data Design', level=1)
    doc.add_paragraph(
        'Data Flow Description: Data flows through the system in a unidirectional pattern from client to server. '
        'User interactions in the React frontend generate data objects that are validated using Zod schemas '
        'and React Hook Form, then sent via Axios HTTP requests to the Express backend API. The backend '
        'validates requests, executes business logic, and performs CRUD operations against the PostgreSQL '
        'database using the pg (node-postgres) driver. Responses flow back to the frontend for rendering. '
        'When offline, data is intercepted by the Sync Service, stored in IndexedDB, and replayed to the '
        'API server upon reconnection.'
    )

    doc.add_heading('7.1 Data Entities', level=2)
    make_table(doc, ['Entity Name', 'Description'], [
        ['User', 'System users including field officers, supervisors, and managers. Stores authentication credentials (email, bcrypt password hash), role assignment, team affiliation, profile photo, language preference, and account status (active/inactive).'],
        ['Citizen', 'Registered citizens for National ID programs. Contains biographic data (name, date of birth, gender), contact information (phone, email), address hierarchy (region, district, village), occupation, marital status, document type and number, biometric collection status, and GPS coordinates of registration.'],
        ['Report', 'Field activity reports submitted by officers. Includes title, description, officer reference, timestamp, photo attachments, GPS location, and sync status (pending/synced).'],
        ['Task', 'Tasks assigned by supervisors to field officers. Contains title, description, assignee, creator, priority level, due date, and completion status.'],
        ['Permission', 'Permission requests (travel, equipment) submitted by officers for supervisor approval. Stores request type, reason, date range, requester, approver, and approval status (pending/approved/rejected).'],
        ['Alert', 'Internal messaging and notification records. Stores sender, recipient, subject, message body, read status, and timestamp.'],
        ['Audit', 'Audit trail records logging all significant system actions. Captures user, action type, target entity, timestamp, and metadata for compliance and accountability.'],
        ['ScreenTime', 'Application usage tracking records. Stores user reference, session start/end times, duration, and device information.'],
        ['Verification', 'Periodic identity verification records for field officers. Tracks verification prompts, responses, timestamps, and compliance status.'],
        ['Location', 'Geographic location reference data (regions, districts, villages) used for cascading selection in citizen registration and reporting.'],
        ['SupervisorReport', 'Summary reports generated by supervisors covering team performance metrics, activity summaries, and operational statistics for a defined period.'],
    ])

    doc.add_paragraph()
    doc.add_heading('7.2 Entity Relationship Diagram', level=2)
    er_path = os.path.join(OUT_DIR, 'diagram_er.png')
    if os.path.exists(er_path):
        doc.add_picture(er_path, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        'Figure 2: Data Entity Relationships \u2014 All entities link to the central users table via foreign keys.'
    ).runs[0].font.size = Pt(9)

    doc.add_paragraph()
    doc.add_heading('7.3 Data Flow Descriptions', level=2)
    make_table(doc, ['Data Flow', 'Description'], [
        ['Citizen Registration Flow', 'Field Officer enters citizen data in the React form (validated by Zod). If online, data is sent via POST /api/citizens to the Express backend, which validates and inserts into the PostgreSQL citizens table. If offline, Dexie.js stores the record in IndexedDB and the Sync Service queues it. Upon reconnection, the Sync Service replays the POST request and updates the local sync status.'],
        ['Authentication Flow', 'User enters email and password on the Login page. The frontend sends a POST /api/auth/login request. The backend retrieves the user by email, compares the bcrypt hash, generates a JWT token, and returns the user object. The frontend stores the token and user context in React state and provides it with all subsequent API requests via Authorization header.'],
        ['Report Submission Flow', 'Officer creates a report with text and optional photo attachments via Multer. If online, the request flows to POST /api/reports, where the backend stores the report in PostgreSQL and saves uploaded files to the uploads Docker volume. If offline, data is stored locally and synced later.'],
        ['Task Assignment Flow', 'Supervisor creates a task via the Tasks interface. POST /api/tasks is sent to the backend, which inserts the task record and triggers an alert notification to the assigned field officer via the alerts table. The officer receives an alert and can view/update the task.'],
        ['Sync / Offline Flow', 'When offline, all write operations are intercepted by SyncService.js and stored in IndexedDB with status "pending". The NetworkStatus hook detects connectivity restoration and triggers SyncService to replay all pending operations against the respective API endpoints. A sync status indicator provides real-time feedback.'],
        ['Audit Logging Flow', 'Every significant API action (create, update, delete) is logged by the backend audit middleware. The audit controller inserts a record into the audit table with user ID, action type, entity type, entity ID, timestamp, and metadata. Managers can query the audit log for compliance review.'],
        ['Email Notification Flow', 'When trigger events occur (e.g., task assignment, permission approval), the backend uses Nodemailer to send email notifications via the configured SMTP server. Email templates are generated server-side and sent to the relevant recipients.'],
    ])

    # ── Diagrams ──
    doc.add_paragraph()
    doc.add_heading('7.4 Citizen Registration Flow Diagram', level=2)
    citizen_path = os.path.join(OUT_DIR, 'diagram_citizen_flow.png')
    if os.path.exists(citizen_path):
        doc.add_picture(citizen_path, width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        'Figure 3: Citizen Registration Data Flow \u2014 Online path (green) through Express API to PostgreSQL, '
        'and offline path (orange) through Dexie.js/IndexedDB with pending sync queue and automatic reconnection.'
    ).runs[0].font.size = Pt(9)

    doc.add_paragraph()
    doc.add_heading('7.5 Authentication Flow Diagram', level=2)
    auth_path = os.path.join(OUT_DIR, 'diagram_auth_flow.png')
    if os.path.exists(auth_path):
        doc.add_picture(auth_path, width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        'Figure 4: Authentication Data Flow \u2014 Step-by-step process from login through bcrypt verification, '
        'JWT token generation, and subsequent authenticated API requests.'
    ).runs[0].font.size = Pt(9)

    doc.add_paragraph()
    doc.add_heading('7.6 Offline Sync Flow Diagram', level=2)
    sync_path = os.path.join(OUT_DIR, 'diagram_sync_flow.png')
    if os.path.exists(sync_path):
        doc.add_picture(sync_path, width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        'Figure 5: Offline Sync Data Flow \u2014 Online path (green) vs offline path (orange) with '
        'reconnection-triggered replay of pending operations.'
    ).runs[0].font.size = Pt(9)

    # ═══════════════════════════════════════════════════════
    # 8. INTERFACES
    # ═══════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading('8. Interfaces', level=1)

    doc.add_heading('8.1 External Interfaces', level=2)
    make_table(doc, ['Interface', 'Description'], [
        ['PostgreSQL Database (TCP)', 'TCP socket connection on port 5432. The backend connects to PostgreSQL using the node-postgres (pg) driver with connection pooling. Used for all persistent data storage and retrieval operations.'],
        ['SMTP Email Server', 'SMTP/TLS connection configured via EMAIL_USER and EMAIL_PASS environment variables. Used by Nodemailer to send password reset notifications, task alerts, and permission status updates.'],
        ['Browser IndexedDB API', 'Client-side browser API accessed via Dexie.js. Used for offline data storage, PWA caching (via vite-plugin-pwa service worker), and client-side state persistence.'],
        ['Browser Geolocation API', 'Client-side browser API for GPS coordinate capture. Used during citizen registration to record location data. Requires user permission.'],
        ['Browser MediaDevices API', 'Client-side browser API for camera access. Used for profile photo capture and citizen document scanning on mobile devices.'],
        ['Docker Engine API', 'Docker Compose orchestrates container lifecycle (build, start, stop, health checks) via the Docker daemon. Used for deployment, scaling, and service management.'],
    ])

    doc.add_paragraph()
    doc.add_heading('8.2 Internal Interfaces', level=2)
    make_table(doc, ['Interface', 'Description'], [
        ['REST API (HTTP/JSON)',
         'The primary internal interface between frontend and backend. All communication uses JSON-formatted HTTP requests and responses. '
         'Endpoints are organized by resource: /api/auth, /api/users, /api/citizens, /api/reports, /api/tasks, '
         '/api/permissions, /api/alerts, /api/audit, /api/screentime, /api/verification, '
         '/api/locations, /api/supervisor-reports, /api/sync. '
         'All endpoints (except login) require JWT Bearer token authentication.'],
        ['Frontend Route Navigation',
         'React Router DOM v6 handles client-side routing. Routes are defined in App.jsx and mapped to components. '
         'Protected routes check authentication state via AuthContext. Role-based route guards restrict '
         'access to supervisor and manager pages.'],
        ['MVC Layer (Backend)',
         'The backend follows Model-View-Controller architecture. Routes (entry points) delegate to Controllers '
         '(business logic), which interact with Models (database queries) and may return Views (API response formatting). '
         'Controllers: auth, user, citizen, report, task, permission, alert, audit, screenTime, '
         'verification, location, supervisorReport, sync.'],
        ['Service Layer (Frontend)',
         'Frontend service modules encapsulate API communication and business logic. SyncService.js manages '
         'offline queue and synchronization. database.js manages IndexedDB operations via Dexie.js. '
         'AuthContext.jsx provides global authentication state to the React component tree.'],
        ['Configuration Layer',
         'Environment-based configuration managed through .env files. Backend config modules: env.ts (environment variables), '
         'db.ts (database connection), mail.ts (SMTP configuration), upload.ts (file storage paths). '
         'Frontend constants and validators in utils/.'],
    ])

    # ═══════════════════════════════════════════════════════
    # 9. SECURITY CONSIDERATIONS
    # ═══════════════════════════════════════════════════════
    doc.add_heading('9. Security Considerations', level=1)
    make_table(doc, ['Security Area', 'Implementation'], [
        ['Authentication', 'JWT-based stateless authentication with bcrypt password hashing (10 salt rounds). Tokens are issued upon successful login and validated on every protected API endpoint. The JWT_SECRET is configured via environment variables and should be a strong, unique string.'],
        ['Authorization', "Role-based access control (RBAC) enforced at both API and UI levels. Backend middleware checks user roles before processing requests. Frontend conditionally renders navigation and features based on the user's role (Manager, Supervisor, Field Officer)."],
        ['Password Security', 'Passwords are hashed using bcrypt with 10 salt rounds before storage. Plaintext passwords are never stored or logged. The change password endpoint verifies the current password before allowing a new one to be set.'],
        ['Data in Transit', 'All API communication uses HTTP in development. For production, HTTPS should be configured via a reverse proxy (e.g., Nginx with SSL/TLS certificates) or a load balancer.'],
        ['Data at Rest', 'PostgreSQL data is stored in Docker volumes with filesystem-level permissions. Sensitive configuration (database credentials, JWT secret, email credentials) is managed through environment variables, not hardcoded in source.'],
        ['Input Validation', 'Server-side validation on all API endpoints using input sanitization. Frontend form validation using Zod schemas and React Hook Form to prevent malformed data submission.'],
        ['CORS Configuration', 'Cross-Origin Resource Sharing (CORS) middleware restricts API access to authorized origins. Configured in the Express backend via the cors package.'],
        ['File Upload Security', 'File uploads handled by Multer with configured storage limits and destination paths. Upload directory is persistent via Docker volumes and should not be publicly accessible.'],
        ['Audit Trail', 'Comprehensive audit logging records all significant system actions including user actions, data modifications, and authentication events. Audit logs are append-only and visible to managers.'],
        ['Session Management', 'JWT tokens are used for session management. The frontend stores tokens in application state (not localStorage in production). Tokens should have appropriate expiration times configured.'],
    ])

    # ═══════════════════════════════════════════════════════
    # 10. PERFORMANCE REQUIREMENTS
    # ═══════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_heading('10. Performance Requirements', level=1)
    add_bullet_list(doc, [
        'API Response Time: All API endpoints shall respond within 500ms under normal load conditions (measured at the 95th percentile).',
        'Page Load Time: The initial page load (including React bundle, CSS, and assets) shall complete within 3 seconds on a standard broadband connection (10 Mbps).',
        'Offline Transition: The application shall transition to offline mode within 2 seconds of network disconnection detection, with no data loss.',
        'Sync Throughput: The Sync Service shall be capable of syncing 50 or more pending records within 60 seconds upon reconnection, subject to server capacity.',
        'Database Query Performance: All standard CRUD queries shall execute within 100ms. Complex analytics queries may take up to 2 seconds with proper indexing.',
        'Concurrent Users: The system shall support a minimum of 100 concurrent users without degradation in response times, using the default Docker Compose deployment.',
        'File Upload: Photo uploads up to 5 MB shall complete within 10 seconds on a standard broadband connection.',
        'Memory Usage: The backend Node.js process shall not exceed 512 MB of RAM under normal operating load.',
        'Storage Capacity: The system shall support at least 10,000 citizen registrations and 5,000 reports before requiring database maintenance or archival.',
        'Availability: The system shall maintain 99.5% uptime during operational hours (8 AM - 5 PM) with planned maintenance windows excluded.',
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # 11. GLOSSARY
    # ═══════════════════════════════════════════════════════
    doc.add_heading('11. Glossary', level=1)
    make_table(doc, ['Term', 'Definition'], [
        ['API', 'Application Programming Interface \u2014 a set of protocols and tools for building software applications.'],
        ['bcrypt', 'A password hashing library designed to be slow and resource-intensive to resist brute-force attacks.'],
        ['CRUD', 'Create, Read, Update, Delete \u2014 the four basic operations of persistent storage.'],
        ['CORS', 'Cross-Origin Resource Sharing \u2014 a mechanism that allows restricted resources on a web page to be requested from another domain.'],
        ['Dexie.js', 'A wrapper library for IndexedDB that provides a simplified, Promise-based API for client-side database operations.'],
        ['Docker Compose', 'A tool for defining and running multi-container Docker applications using a YAML configuration file.'],
        ['IndexedDB', 'A low-level browser API for client-side storage of significant amounts of structured data, including files and blobs.'],
        ['i18next', 'An internationalization framework for JavaScript that provides translation, interpolation, and language detection.'],
        ['JWT', 'JSON Web Token \u2014 a compact, URL-safe means of representing claims to be transferred between two parties.'],
        ['MVC', 'Model-View-Controller \u2014 a software design pattern that separates an application into three interconnected components.'],
        ['Multer', 'A Node.js middleware for handling multipart/form-data, primarily used for file uploads.'],
        ['Nodemailer', 'A Node.js module for sending emails via SMTP, SMTPD, SES, and other transport methods.'],
        ['Nginx', 'A high-performance web server and reverse proxy used for serving static files and load balancing.'],
        ['PWA', 'Progressive Web App \u2014 a type of web application that provides native app-like features including offline access and home screen installation.'],
        ['RBAC', 'Role-Based Access Control \u2014 a method of regulating access to resources based on user roles.'],
        ['REST', 'Representational State Transfer \u2014 an architectural style for designing networked applications.'],
        ['SPA', 'Single Page Application \u2014 a web application that loads a single HTML page and dynamically updates content.'],
        ['Sync Service', 'A client-side module that queues offline operations and replays them against the API server upon reconnection.'],
        ['Vite', 'A fast frontend build tool that provides instant hot module replacement and optimized production builds.'],
        ['Zod', 'A TypeScript-first schema declaration and validation library used for form and API input validation.'],
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # 12. APPENDICES
    # ═══════════════════════════════════════════════════════
    doc.add_heading('12. Appendices', level=1)

    doc.add_heading('A. Backend API Routes', level=2)
    doc.add_paragraph('The following route modules are defined in the backend Express application:')
    make_table(doc, ['Route File', 'Base Path', 'Description'], [
        ['auth.routes.ts', '/api/auth', 'Login, password change'],
        ['user.routes.ts', '/api/users', 'User CRUD, profile management'],
        ['citizen.routes.ts', '/api/citizens', 'Citizen registration and queries'],
        ['report.routes.ts', '/api/reports', 'Report creation and retrieval'],
        ['task.routes.ts', '/api/tasks', 'Task assignment and tracking'],
        ['permission.routes.ts', '/api/permissions', 'Permission request/approval workflow'],
        ['alert.routes.ts', '/api/alerts', 'Alert/notification messaging'],
        ['audit.routes.ts', '/api/audit', 'Audit log queries'],
        ['screenTime.routes.ts', '/api/screentime', 'Application usage tracking'],
        ['verification.routes.ts', '/api/verification', 'Identity verification prompts'],
        ['location.routes.ts', '/api/locations', 'Geographic reference data'],
        ['supervisorReport.routes.ts', '/api/supervisor-reports', 'Supervisor report generation'],
        ['sync.routes.ts', '/api/sync', 'Offline data synchronization'],
    ])

    doc.add_paragraph()
    doc.add_heading('B. Docker Compose Services', level=2)
    make_table(doc, ['Service', 'Image', 'Role', 'Ports', 'Volumes', 'Notes'], [
        ['db', 'postgres:16-alpine', 'PostgreSQL database', '5432 (internal)', 'pgdata volume', 'Health check: pg_isready'],
        ['backend', 'Custom build (./backend)', 'Express API server', '5001:5000', 'uploads volume', 'Depends on: db (healthy)'],
        ['frontend', 'Custom build (./frontend)', 'React SPA + Nginx', '30001:80', 'None', 'Depends on: backend'],
    ])

    doc.add_paragraph()
    doc.add_heading('C. Frontend Technology Stack', level=2)
    make_table(doc, ['Technology', 'Purpose'], [
        ['React 18.2', 'UI Component Library'],
        ['Vite 5.0', 'Build Tool & Dev Server'],
        ['React Router DOM 6.30', 'Client-side Routing'],
        ['Tailwind CSS 3.4', 'Utility-first CSS Framework'],
        ['Axios 1.18', 'HTTP Client'],
        ['Dexie.js 4.4', 'IndexedDB Wrapper (Offline Storage)'],
        ['i18next 23.16', 'Internationalization'],
        ['React Hook Form 7.80', 'Form State Management'],
        ['Zod 3.25', 'Schema Validation'],
        ['Recharts 3.9', 'Charting & Data Visualization'],
        ['Date-fns 2.30', 'Date Formatting & Manipulation'],
        ['React Hot Toast 2.6', 'Toast Notifications'],
        ['Vite Plugin PWA 0.17', 'Progressive Web App Support'],
    ])

    doc.add_paragraph()
    doc.add_heading('D. Backend Technology Stack', level=2)
    make_table(doc, ['Technology', 'Purpose'], [
        ['Node.js (LTS)', 'JavaScript Runtime'],
        ['Express 5.2', 'Web Framework'],
        ['TypeScript 5.6', 'Type-safe JavaScript'],
        ['pg 8.22', 'PostgreSQL Client Driver'],
        ['bcrypt 6.0', 'Password Hashing'],
        ['dotenv 17.4', 'Environment Variable Management'],
        ['cors 2.8', 'CORS Middleware'],
        ['multer 2.0', 'File Upload Handling'],
        ['nodemailer 9.0', 'Email Sending'],
        ['tsx 4.19', 'TypeScript Execution (Dev)'],
    ])

    doc.add_paragraph()
    doc.add_heading('E. Environment Configuration Reference', level=2)
    make_table(doc, ['Variable', 'Default', 'Description'], [
        ['PORT', '5000', 'Backend server listen port'],
        ['DB_USER', 'postgres', 'PostgreSQL authentication username'],
        ['DB_PASSWORD', '(required)', 'PostgreSQL authentication password'],
        ['DB_HOST', 'db', 'Database hostname (use "db" within Docker network)'],
        ['DB_PORT', '5432', 'PostgreSQL port'],
        ['DB_NAME', 'fieldsync_db', 'Database name'],
        ['JWT_SECRET', '(required)', 'Secret key for JWT token signing and verification'],
        ['EMAIL_USER', '(required)', 'SMTP email address for outgoing notifications'],
        ['EMAIL_PASS', '(required)', 'SMTP email password or application-specific password'],
    ])

    # ── SAVE ──
    output_path = os.path.join(OUT_DIR, 'FieldSync_System_Design_Document.docx')
    doc.save(output_path)
    print(f'SDD saved to: {output_path}')
    return output_path


if __name__ == '__main__':
    print('Generating diagrams...')
    print(f'  [OK] {create_architecture_diagram()}')
    print(f'  [OK] {create_citizen_flow_diagram()}')
    print(f'  [OK] {create_auth_flow_diagram()}')
    print(f'  [OK] {create_sync_flow_diagram()}')
    print(f'  [OK] {create_er_diagram()}')

    print('\nGenerating SDD document...')
    path = generate_sdd()
    print(f'\nDone! {path}')
