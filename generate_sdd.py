from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# --- Style Configuration ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(30, 58, 95)
    hs.font.name = 'Calibri'

NAVY = '1E3A5F'
LIGHT_BLUE = '4FC3F7'

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells_data, bold=False, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        if bold or header:
            run.bold = True
        if header:
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_shading(cell, NAVY)
    return row

def make_table(doc, headers, rows):
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, headers, header=True)
    for row_data in rows:
        add_table_row(table, row_data)
    return table

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(50, 50, 50)

# =============================================
# COVER PAGE
# =============================================
for _ in range(5):
    doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('FieldSync')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(30, 58, 95)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('System Design Document')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(79, 195, 247)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Architecture, Components, Interfaces, and Data Flows')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f'Version 1.0.0\n{datetime.date.today().strftime("%B %Y")}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 116, 139)

doc.add_page_break()

# =============================================
# TABLE OF CONTENTS
# =============================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Overview',
    '2. System Objectives',
    '3. Scope',
    '4. Assumptions and Constraints',
    '5. Stakeholders',
    '6. System Architecture',
    '7. Data Design',
    '8. Interfaces',
    '9. Security Considerations',
    '10. Performance Requirements',
    '11. Glossary',
    '12. Appendices',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(4)

doc.add_page_break()

# =============================================
# 1. OVERVIEW
# =============================================
doc.add_heading('1. Overview', level=1)

overview_rows = [
    ['System Name', 'FieldSync'],
    ['Prepared By', 'FieldSync Development Team'],
    ['Date', datetime.date.today().strftime('%Y-%m-%d')],
    ['Version', '1.0.0'],
    ['Description', 'A comprehensive field workforce management and citizen registration platform designed to streamline government field operations including National ID registration, reporting, task management, team oversight, and analytics.'],
]
make_table(doc, ['Field', 'Details'], overview_rows)

# =============================================
# 2. SYSTEM OBJECTIVES
# =============================================
doc.add_paragraph()
doc.add_heading('2. System Objectives', level=1)

objectives = [
    'Enable field officers to register citizens for National ID programs accurately with GPS-verified location data and document capture.',
    'Provide offline-first architecture that ensures uninterrupted field operations regardless of internet connectivity, with automatic data synchronization upon reconnection.',
    'Establish a hierarchical role-based system (Field Officer, Supervisor, Manager) with granular access control for secure and efficient operations.',
    'Facilitate real-time task assignment, tracking, and completion workflows between supervisors and field officers.',
    'Deliver comprehensive analytics, audit trails, and reporting capabilities for managerial decision-making and regulatory compliance.',
    'Support multi-language accessibility (English, Amharic, Tigrinya, Oromo) to serve a diverse workforce.',
    'Ensure data integrity, security, and availability through encrypted storage, authentication, and automated backup mechanisms.',
    'Reduce paper-based processes and manual data entry by digitizing citizen registration and field reporting workflows.',
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

# =============================================
# 3. SCOPE
# =============================================
doc.add_heading('3. Scope', level=1)

doc.add_heading('3.1 In Scope', level=2)
in_scope = [
    'Citizen registration for National ID with biographic data, document types (National ID, Birth Certificate, Passport), GPS capture, and biometric collection status.',
    'User authentication and role-based authorization (Manager, Supervisor, Field Officer) with JWT tokens and bcrypt password hashing.',
    'Field reporting system with offline creation, photo/document attachment support, and automatic server synchronization.',
    'Task management workflow including creation, assignment, prioritization, due dates, and completion tracking.',
    'Team management for supervisors to monitor and manage field officer activities.',
    'Permission request and approval workflow for leaves, travel, and equipment requests.',
    'Supervisor report generation summarizing team performance and activities.',
    'Screen time tracking and identity verification for field officer monitoring.',
    'Alert and messaging system for internal notifications and team communication.',
    'Manager analytics dashboard with trend charts, registration statistics, and performance metrics.',
    'Audit logging of all system actions for accountability and compliance.',
    'Progressive Web App (PWA) support for mobile installation and offline access.',
    'Multi-language support with English, Amharic, Tigrinya, and Oromo translations.',
    'Dark and light theme toggle for user interface customization.',
    'Docker-based deployment with automated database migration and service orchestration.',
]
for item in in_scope:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.2 Out of Scope', level=2)
out_scope = [
    'Biometric data capture and storage (only collection status is tracked; actual biometric processing is handled by external systems).',
    'Payment processing or financial transaction management.',
    'Integration with external national ID issuance systems (data is prepared for export only).',
    'Mobile native applications (iOS/Android) -- the system is web-based with PWA support.',
    'Video conferencing or real-time video streaming capabilities.',
]
for item in out_scope:
    doc.add_paragraph(item, style='List Bullet')

# =============================================
# 4. ASSUMPTIONS AND CONSTRAINTS
# =============================================
doc.add_heading('4. Assumptions and Constraints', level=1)

doc.add_heading('4.1 Assumptions', level=2)
assumptions = [
    'Docker and Docker Compose are available and properly installed on the deployment server.',
    'The target deployment environment has stable power supply and reasonable network connectivity for initial setup.',
    'Field officers have access to modern web browsers (Chrome 90+, Firefox 88+, Edge 90+, or Safari 14+) on their devices.',
    'GPS and camera hardware is available on field officer devices for location capture and photo uploads.',
    'Administrators have basic command-line proficiency for Docker-based deployment and configuration.',
    'The PostgreSQL database server will be managed within Docker containers and does not require an external managed database service.',
    'Email services (SMTP) are available for notification delivery via Nodemailer.',
]
for item in assumptions:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.2 Constraints', level=2)
constraints = [
    'The system requires Docker and Docker Compose for deployment; no native bare-metal installation path is provided.',
    'Offline data is stored in the browser\'s IndexedDB, which has device-dependent storage limits (typically 50 MB to several GB depending on browser and OS).',
    'All API communication uses HTTP/HTTPS; no WebSocket or real-time push notification infrastructure is currently in place.',
    'The frontend is served via Nginx in production; custom server-side rendering is not supported.',
    'File uploads (photos, documents) are limited to server storage capacity, which is backed by Docker volumes.',
    'Multi-language support covers the four specified languages; adding new languages requires manual translation updates.',
    'The system does not support multi-tenancy; each deployment serves a single organizational unit.',
]
for item in constraints:
    doc.add_paragraph(item, style='List Bullet')

# =============================================
# 5. STAKEHOLDERS
# =============================================
doc.add_heading('5. Stakeholders', level=1)

stakeholders = [
    ['Project Sponsor', 'Program Director', 'sponsor@fieldsync.gov'],
    ['Project Manager', 'Delivery Lead', 'pm@fieldsync.gov'],
    ['Lead Developer', 'Backend & Architecture', 'lead-dev@fieldsync.gov'],
    ['Frontend Developer', 'UI/UX Implementation', 'frontend-dev@fieldsync.gov'],
    ['Backend Developer', 'API & Database', 'backend-dev@fieldsync.gov'],
    ['QA Engineer', 'Testing & Quality Assurance', 'qa@fieldsync.gov'],
    ['DevOps Engineer', 'Deployment & Infrastructure', 'devops@fieldsync.gov'],
    ['Field Officers', 'End Users (Registration & Reporting)', 'field-team@fieldsync.gov'],
    ['Supervisors', 'End Users (Team Management)', 'supervisors@fieldsync.gov'],
    ['Managers', 'End Users (Analytics & Administration)', 'managers@fieldsync.gov'],
    ['Database Administrator', 'Database Management & Backup', 'dba@fieldsync.gov'],
    ['Security Officer', 'Security & Compliance Review', 'security@fieldsync.gov'],
]
make_table(doc, ['Name / Role', 'Responsibility', 'Contact'], stakeholders)

# =============================================
# 6. SYSTEM ARCHITECTURE
# =============================================
doc.add_paragraph()
doc.add_heading('6. System Architecture', level=1)

doc.add_paragraph(
    'Overview: FieldSync follows a three-tier client-server architecture deployed via Docker Compose. '
    'The system comprises a React-based single-page application (SPA) frontend served by Nginx, '
    'a Node.js/Express REST API backend written in TypeScript, and a PostgreSQL 16 relational database. '
    'The frontend communicates with the backend exclusively through RESTful HTTP APIs. '
    'Offline-first capability is achieved via Dexie.js (IndexedDB) on the client side with an '
    'automatic synchronization layer that resolves conflicts upon reconnection.'
)

doc.add_heading('6.1 Components', level=2)

components = [
    ['Frontend (React + Vite)',
     'A responsive Single Page Application built with React 18, Vite 5, and Tailwind CSS 3.4. '
     'Served by Nginx in production. Implements offline-first data storage using Dexie.js (IndexedDB wrapper), '
     'role-based UI rendering, multi-language support via i18next, form validation with React Hook Form + Zod, '
     'charts with Recharts, and PWA capabilities via vite-plugin-pwa.'],
    ['Backend API (Express + TypeScript)',
     'A RESTful API server built with Express 5 and TypeScript, following MVC architecture. '
     'Handles authentication (bcrypt + JWT), request validation, business logic, file uploads (Multer), '
     'and email notifications (Nodemailer). Runs on Node.js and exposes endpoints for auth, users, citizens, '
     'reports, tasks, permissions, alerts, audits, screen time, verification, locations, '
     'and supervisor reports.'],
    ['Database (PostgreSQL 16)',
     'Relational database using PostgreSQL 16 Alpine Docker image. Stores all persistent data including '
     'user accounts, citizen registrations, reports, tasks, permissions, alerts, audit logs, '
     'screen time data, and verification records. Data is persisted via Docker named volumes.'],
    ['Nginx (Reverse Proxy)',
     'Lightweight web server used to serve the frontend static build and proxy API requests to the backend. '
     'Configured as part of the frontend Docker container.'],
    ['Docker Compose',
     'Container orchestration layer that defines and manages three services (db, backend, frontend) with '
     'health checks, dependency ordering, volume mounts, and network configuration.'],
    ['Sync Service (Client-side)',
     'A client-side synchronization module (SyncService.js) that queues data operations in IndexedDB when '
     'offline and replays them against the backend API when connectivity is restored. Handles conflict resolution '
     'and provides sync status indicators.'],
    ['Offline Storage (Dexie.js / IndexedDB)',
     'Client-side database layer using Dexie.js as a wrapper around the browser\'s IndexedDB API. Stores '
     'citizen registrations, reports, and pending sync operations for offline access.'],
]
make_table(doc, ['Component', 'Description'], components)

doc.add_paragraph()
doc.add_heading('6.2 System Diagram', level=2)

doc.add_paragraph(
    'The system architecture follows a layered pattern with clear separation of concerns:'
)

doc.add_paragraph()
diagram_text = (
    '┌──────────────────────────────────────────────────────────────────┐\n'
    '│                    CLIENT BROWSER (PWA)                         │\n'
    '│  ┌─────────────────┐  ┌──────────────┐  ┌───────────────────┐  │\n'
    '│  │   React SPA      │  │  Dexie.js /  │  │  Sync Service     │  │\n'
    '│  │   (UI / Routing) │  │  IndexedDB   │  │  (Offline Queue)  │  │\n'
    '│  │   i18next        │  │  (Offline    │  │  (Auto Sync)      │  │\n'
    '│  │   Recharts       │  │   Storage)   │  │                   │  │\n'
    '│  └─────────────────┘  └──────────────┘  └───────────────────┘  │\n'
    '└────────────────────────────┬────────────────────────────────────┘\n'
    '                             │  HTTP/REST (JSON)\n'
    '┌────────────────────────────▼────────────────────────────────────┐\n'
    '│                    NGINX REVERSE PROXY                         │\n'
    '│              (Port 80 → Frontend Static / API Proxy)           │\n'
    '└────────────────────────────┬────────────────────────────────────┘\n'
    '                             │\n'
    '┌────────────────────────────▼────────────────────────────────────┐\n'
    '│              NODE.JS / EXPRESS API SERVER                       │\n'
    '│  ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────────────┐  │\n'
    '│  │ Auth     │ │ Citizen  │ │ Report   │ │ Task / Permission│  │\n'
    '│  │ Module   │ │ Module   │ │ Module   │ │ / Alert Module   │  │\n'
    '│  └──────────┘ └──────────┘ └──────────┘ └──────────────────┘  │\n'
    '│  ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────────────┐  │\n'
    '│  │ User     │ │ Audit    │ │ Sync     │ │ File Upload      │  │\n'
    '│  │ Module   │ │ Module   │ │ Module   │ │ (Multer)         │  │\n'
    '│  └──────────┘ └──────────┘ └──────────┘ └──────────────────┘  │\n'
    '│  ┌──────────────────────────────────────────────────────────┐  │\n'
    '│  │              Nodemailer (Email Notifications)            │  │\n'
    '│  └──────────────────────────────────────────────────────────┘  │\n'
    '└────────────────────────────┬────────────────────────────────────┘\n'
    '                             │  TCP (pg)\n'
    '┌────────────────────────────▼────────────────────────────────────┐\n'
    '│              PostgreSQL 16 (ALPINE)                            │\n'
    '│  ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────────────┐  │\n'
    '│  │ users    │ │ citizens │ │ reports  │ │ tasks            │  │\n'
    '│  │ alerts   │ │ audit    │ │ screen_  │ │ permissions      │  │\n'
    '│  │          │ │ verif.   │ │ time     │ │ supervisor_rpts  │  │\n'
    '│  │          │ │          │ │          │ │ locations        │  │\n'
    '│  └──────────┘ └──────────┘ └──────────┘ └──────────────────┘  │\n'
    '└──────────────────────────────────────────────────────────────────┘\n'
)

p = doc.add_paragraph()
run = p.add_run(diagram_text)
run.font.name = 'Consolas'
run.font.size = Pt(7)

doc.add_paragraph()
doc.add_paragraph(
    'Diagram Image/Link: A detailed Visio/Lucidchart version of this diagram should be maintained '
    'at the project documentation repository. The above ASCII representation provides an overview '
    'of the three-tier architecture and component relationships.'
)

# =============================================
# 7. DATA DESIGN
# =============================================
doc.add_heading('7. Data Design', level=1)

doc.add_paragraph(
    'Data Flow Description: Data flows through the system in a unidirectional pattern from client to server. '
    'User interactions in the React frontend generate data objects that are validated using Zod schemas '
    'and React Hook Form, then sent via Axios HTTP requests to the Express backend API. The backend '
    'validates requests, executes business logic, and performs CRUD operations against the PostgreSQL '
    'database using the pg (node-postgres) driver. Responses flow back to the frontend for rendering. '
    'When offline, data is intercepted by the Sync Service, stored in IndexedDB, and replayed to the '
    'API server upon reconnection.'
)

doc.add_heading('7.1 Data Entities', level=2)

entities = [
    ['User', 'System users including field officers, supervisors, and managers. Stores authentication credentials (email, bcrypt password hash), role assignment, team affiliation, profile photo, language preference, and account status (active/inactive).'],
    ['Citizen', 'Registered citizens for National ID programs. Contains biographic data (name, date of birth, gender), contact information (phone, email), address hierarchy (region, district, village), occupation, marital status, document type and number, biometric collection status, and GPS coordinates of registration.'],
    ['Report', 'Field activity reports submitted by officers. Includes title, description, officer reference, timestamp, photo attachments, GPS location, and sync status (pending/synced).'],
    ['Task', 'Tasks assigned by supervisors to field officers. Contains title, description, assignee, creator, priority level, due date, and completion status.'],
    ['Permission', 'Permission requests (travel, equipment) submitted by officers for supervisor approval. Stores request type, reason, date range, requester, approver, and approval status (pending/approved/rejected).'],
    ['Alert', 'Internal messaging and notification records. Stores sender, recipient, subject, message body, read status, and timestamp.'],
    ['Audit', 'Audit trail records logging all significant system actions. Captures user, action type, target entity, timestamp, and metadata for compliance and accountability.'],
    ['ScreenTime', 'Application usage tracking records. Stores user reference, session start/end times, duration, and device information.'],
    ['Verification', 'Periodic identity verification records for field officers. Tracks verification prompts, responses, timestamps, and compliance status.'],
    ['Location', 'Geographic location reference data (regions, districts, villages) used for cascading selection in citizen registration and reporting.'],
    ['SupervisorReport', 'Summary reports generated by supervisors covering team performance metrics, activity summaries, and operational statistics for a defined period.'],
]
make_table(doc, ['Entity Name', 'Description'], entities)

doc.add_paragraph()
doc.add_heading('7.2 Data Flow Diagrams', level=2)

data_flows = [
    ['Citizen Registration Flow',
     'Field Officer enters citizen data in the React form (validated by Zod). If online, data is sent via '
     'POST /api/citizens to the Express backend, which validates and inserts into the PostgreSQL citizens table. '
     'If offline, Dexie.js stores the record in IndexedDB and the Sync Service queues it. Upon reconnection, '
     'the Sync Service replays the POST request and updates the local sync status.'],
    ['Authentication Flow',
     'User enters email and password on the Login page. The frontend sends a POST /api/auth/login request. '
     'The backend retrieves the user by email, compares the bcrypt hash, generates a JWT token, and returns '
     'the user object. The frontend stores the token and user context in React state and provides it with '
     'all subsequent API requests via Authorization header.'],
    ['Report Submission Flow',
     'Officer creates a report with text and optional photo attachments via Multer. If online, the request '
     'flows to POST /api/reports, where the backend stores the report in PostgreSQL and saves uploaded files '
     'to the uploads Docker volume. If offline, data is stored locally and synced later.'],
    ['Task Assignment Flow',
     'Supervisor creates a task via the Tasks interface. POST /api/tasks is sent to the backend, which '
     'inserts the task record and triggers an alert notification to the assigned field officer via the alerts table. '
     'The officer receives an alert and can view/update the task.'],
    ['Sync / Offline Flow',
     'When offline, all write operations are intercepted by SyncService.js and stored in IndexedDB with status "pending". '
     'The NetworkStatus hook detects connectivity restoration and triggers SyncService to replay all pending '
     'operations against the respective API endpoints. A sync status indicator provides real-time feedback.'],
    ['Audit Logging Flow',
     'Every significant API action (create, update, delete) is logged by the backend audit middleware. '
     'The audit controller inserts a record into the audit table with user ID, action type, entity type, '
     'entity ID, timestamp, and metadata. Managers can query the audit log for compliance review.'],
    ['Email Notification Flow',
     'When trigger events occur (e.g., task assignment, permission approval), the backend uses Nodemailer '
     'to send email notifications via the configured SMTP server. Email templates are generated server-side '
     'and sent to the relevant recipients.'],
]
make_table(doc, ['Data Flow', 'Description'], data_flows)

# =============================================
# 8. INTERFACES
# =============================================
doc.add_heading('8. Interfaces', level=1)

doc.add_heading('8.1 External Interfaces', level=2)

external_ifs = [
    ['PostgreSQL Database (TCP)', 'TCP socket connection on port 5432. The backend connects to PostgreSQL using the node-postgres (pg) driver with connection pooling. Used for all persistent data storage and retrieval operations.'],
    ['SMTP Email Server', 'SMTP/TLS connection configured via EMAIL_USER and EMAIL_PASS environment variables. Used by Nodemailer to send password reset notifications, task alerts, and permission status updates.'],
    ['Browser IndexedDB API', 'Client-side browser API accessed via Dexie.js. Used for offline data storage, PWA caching (via vite-plugin-pwa service worker), and client-side state persistence.'],
     ['Browser Geolocation API', 'Client-side browser API for GPS coordinate capture. Used during citizen registration to record location data. Requires user permission.'],
    ['Browser MediaDevices API', 'Client-side browser API for camera access. Used for profile photo capture and citizen document scanning on mobile devices.'],
    ['Docker Engine API', 'Docker Compose orchestrates container lifecycle (build, start, stop, health checks) via the Docker daemon. Used for deployment, scaling, and service management.'],
]
make_table(doc, ['Interface', 'Description'], external_ifs)

doc.add_paragraph()
doc.add_heading('8.2 Internal Interfaces', level=2)

internal_ifs = [
    ['REST API (HTTP/JSON)',
     'The primary internal interface between frontend and backend. All communication uses JSON-formatted HTTP requests and responses. '
     'Endpoints are organized by resource: /api/auth, /api/users, /api/citizens, /api/reports, /api/tasks, '
     '/api/permissions, /api/alerts, /api/audit, /api/screentime, /api/verification, '
     '/api/locations, /api/supervisor-reports, /api/sync. '
     'All endpoints (except login) require JWT Bearer token authentication.'],
    ['Frontend Route Navigation',
     'React Router DOM v6 handles client-side routing. Routes are defined in App.jsx and mapped to components. '
     'Protected routes check authentication state via AuthContext. Role-based route guards restrict '
     'access to supervisor and manager pages.'],
    ['MVC Layer (Backend)',
     'The backend follows Model-View-Controller architecture. Routes (entry points) delegate to Controllers '
     '(business logic), which interact with Models (database queries) and may return Views (API response formatting). '
     'Controllers: auth, user, citizen, report, task, permission, alert, audit, screenTime, '
     'verification, location, supervisorReport, sync.'],
    ['Service Layer (Frontend)',
     'Frontend service modules encapsulate API communication and business logic. SyncService.js manages '
     'offline queue and synchronization. database.js manages IndexedDB operations via Dexie.js. '
     'AuthContext.jsx provides global authentication state to the React component tree.'],
    ['Configuration Layer',
     'Environment-based configuration managed through .env files. Backend config modules: env.ts (environment variables), '
     'db.ts (database connection), mail.ts (SMTP configuration), upload.ts (file storage paths). '
     'Frontend constants and validators in utils/.'],
]
make_table(doc, ['Interface', 'Description'], internal_ifs)

# =============================================
# 9. SECURITY CONSIDERATIONS
# =============================================
doc.add_heading('9. Security Considerations', level=1)

security = [
    ['Authentication',
     'JWT-based stateless authentication with bcrypt password hashing (10 salt rounds). Tokens are issued upon successful login and validated on every protected API endpoint. '
     'The JWT_SECRET is configured via environment variables and should be a strong, unique string.'],
    ['Authorization',
     'Role-based access control (RBAC) enforced at both API and UI levels. Backend middleware checks user roles before '
     'processing requests. Frontend conditionally renders navigation and features based on the user\'s role '
     '(Manager, Supervisor, Field Officer).'],
    ['Password Security',
     'Passwords are hashed using bcrypt with 10 salt rounds before storage. Plaintext passwords are never stored or logged. '
     'The change password endpoint verifies the current password before allowing a new one to be set.'],
    ['Data in Transit',
     'All API communication uses HTTP in development. For production, HTTPS should be configured via a reverse proxy '
     '(e.g., Nginx with SSL/TLS certificates) or a load balancer.'],
    ['Data at Rest',
     'PostgreSQL data is stored in Docker volumes with filesystem-level permissions. Sensitive configuration (database credentials, '
     'JWT secret, email credentials) is managed through environment variables, not hardcoded in source.'],
    ['Input Validation',
     'Server-side validation on all API endpoints using input sanitization. Frontend form validation using Zod schemas '
     'and React Hook Form to prevent malformed data submission.'],
    ['CORS Configuration',
     'Cross-Origin Resource Sharing (CORS) middleware restricts API access to authorized origins. '
     'Configured in the Express backend via the cors package.'],
    ['File Upload Security',
     'File uploads handled by Multer with configured storage limits and destination paths. '
     'Upload directory is persistent via Docker volumes and should not be publicly accessible.'],
    ['Audit Trail',
     'Comprehensive audit logging records all significant system actions including user actions, '
     'data modifications, and authentication events. Audit logs are append-only and visible to managers.'],
    ['Session Management',
     'JWT tokens are used for session management. The frontend stores tokens in application state '
     '(not localStorage in production). Tokens should have appropriate expiration times configured.'],
]
make_table(doc, ['Security Area', 'Implementation'], security)

# =============================================
# 10. PERFORMANCE REQUIREMENTS
# =============================================
doc.add_paragraph()
doc.add_heading('10. Performance Requirements', level=1)

perf_reqs = [
    'API Response Time: All API endpoints shall respond within 500ms under normal load conditions (measured at the 95th percentile).',
    'Page Load Time: The initial page load (including React bundle, CSS, and assets) shall complete within 3 seconds on a standard broadband connection (10 Mbps).',
    'Offline Transition: The application shall transition to offline mode within 2 seconds of network disconnection detection, with no data loss.',
    'Sync Throughput: The Sync Service shall be capable of syncing 50 or more pending records within 60 seconds upon reconnection, subject to server capacity.',
    'Database Query Performance: All standard CRUD queries shall execute within 100ms. Complex analytics queries may take up to 2 seconds with proper indexing.',
    'Concurrent Users: The system shall support a minimum of 100 concurrent users without degradation in response times, using the default Docker Compose deployment.',
    'File Upload: Photo uploads up to 5 MB shall complete within 10 seconds on a standard broadband connection.',
    'Memory Usage: The backend Node.js process shall not exceed 512 MB of RAM under normal operating load.',
    'Storage Capacity: The system shall support at least 10,000 citizen registrations and 5,000 reports before requiring database maintenance or archival.',
    'Availability: The system shall maintain 99.5% uptime during operational hours (8 AM - 5 PM) with planned maintenance windows excluded.',
]
for req in perf_reqs:
    doc.add_paragraph(req, style='List Bullet')

# =============================================
# 11. GLOSSARY
# =============================================
doc.add_heading('11. Glossary', level=1)

glossary = [
    ['API', 'Application Programming Interface -- a set of protocols and tools for building software applications.'],
    ['bcrypt', 'A password hashing library designed to be slow and resource-intensive to resist brute-force attacks.'],
    ['CRUD', 'Create, Read, Update, Delete -- the four basic operations of persistent storage.'],
    ['CORS', 'Cross-Origin Resource Sharing -- a mechanism that allows restricted resources on a web page to be requested from another domain.'],
    ['Dexie.js', 'A wrapper library for IndexedDB that provides a simplified, Promise-based API for client-side database operations.'],
    ['Docker Compose', 'A tool for defining and running multi-container Docker applications using a YAML configuration file.'],
    ['IndexedDB', 'A low-level browser API for client-side storage of significant amounts of structured data, including files and blobs.'],
    ['i18next', 'An internationalization framework for JavaScript that provides translation, interpolation, and language detection.'],
    ['JWT', 'JSON Web Token -- a compact, URL-safe means of representing claims to be transferred between two parties.'],
    ['MVC', 'Model-View-Controller -- a software design pattern that separates an application into three interconnected components.'],
    ['Multer', 'A Node.js middleware for handling multipart/form-data, primarily used for file uploads.'],
    ['Nodemailer', 'A Node.js module for sending emails via SMTP, SMTPD, SES, and other transport methods.'],
    ['Nginx', 'A high-performance web server and reverse proxy used for serving static files and load balancing.'],
    ['PWA', 'Progressive Web App -- a type of web application that provides native app-like features including offline access and home screen installation.'],
    ['RBAC', 'Role-Based Access Control -- a method of regulating access to resources based on user roles.'],
    ['REST', 'Representational State Transfer -- an architectural style for designing networked applications.'],
    ['SPA', 'Single Page Application -- a web application that loads a single HTML page and dynamically updates content.'],
    ['Sync Service', 'A client-side module that queues offline operations and replays them against the API server upon reconnection.'],
    ['Vite', 'A fast frontend build tool that provides instant hot module replacement and optimized production builds.'],
    ['Zod', 'A TypeScript-first schema declaration and validation library used for form and API input validation.'],
]
make_table(doc, ['Term', 'Definition'], glossary)

# =============================================
# 12. APPENDICES
# =============================================
doc.add_paragraph()
doc.add_heading('12. Appendices', level=1)

doc.add_heading('A. Backend API Routes', level=2)
doc.add_paragraph('The following route modules are defined in the backend Express application:')

routes = [
    ['auth.routes.ts', '/api/auth', 'Login, password change'],
    ['user.routes.ts', '/api/users', 'User CRUD, profile management'],
    ['citizen.routes.ts', '/api/citizens', 'Citizen registration and queries'],
    ['report.routes.ts', '/api/reports', 'Report creation and retrieval'],
    ['task.routes.ts', '/api/tasks', 'Task assignment and tracking'],
    ['permission.routes.ts', '/api/permissions', 'Permission request/approval workflow'],
    ['alert.routes.ts', '/api/alerts', 'Alert/notification messaging'],
    ['audit.routes.ts', '/api/audit', 'Audit log queries'],
    ['screenTime.routes.ts', '/api/screentime', 'Application usage tracking'],
    ['verification.routes.ts', '/api/verification', 'Identity verification prompts'],
    ['location.routes.ts', '/api/locations', 'Geographic reference data'],
    ['supervisorReport.routes.ts', '/api/supervisor-reports', 'Supervisor report generation'],
    ['sync.routes.ts', '/api/sync', 'Offline data synchronization'],
]
make_table(doc, ['Route File', 'Base Path', 'Description'], routes)

doc.add_paragraph()
doc.add_heading('B. Docker Compose Services', level=2)

docker_services = [
    ['db', 'postgres:16-alpine', 'PostgreSQL database', '5432 (internal)', 'pgdata volume', 'Health check: pg_isready'],
    ['backend', 'Custom build (./backend)', 'Express API server', '5001:5000', 'uploads volume', 'Depends on: db (healthy)'],
    ['frontend', 'Custom build (./frontend)', 'React SPA + Nginx', '30001:80', 'None', 'Depends on: backend'],
]
make_table(doc, ['Service', 'Image', 'Role', 'Ports', 'Volumes', 'Notes'], docker_services)

doc.add_paragraph()
doc.add_heading('C. Frontend Technology Stack', level=2)

fe_stack = [
    ['React 18.2', 'UI Component Library'],
    ['Vite 5.0', 'Build Tool & Dev Server'],
    ['React Router DOM 6.30', 'Client-side Routing'],
    ['Tailwind CSS 3.4', 'Utility-first CSS Framework'],
    ['Axios 1.18', 'HTTP Client'],
    ['Dexie.js 4.4', 'IndexedDB Wrapper (Offline Storage)'],
    ['i18next 23.16', 'Internationalization'],
    ['React Hook Form 7.80', 'Form State Management'],
    ['Zod 3.25', 'Schema Validation'],
    ['Recharts 3.9', 'Charting & Data Visualization'],
    ['Date-fns 2.30', 'Date Formatting & Manipulation'],
    ['React Hot Toast 2.6', 'Toast Notifications'],
    ['Vite Plugin PWA 0.17', 'Progressive Web App Support'],
]
make_table(doc, ['Technology', 'Purpose'], fe_stack)

doc.add_paragraph()
doc.add_heading('D. Backend Technology Stack', level=2)

be_stack = [
    ['Node.js (LTS)', 'JavaScript Runtime'],
    ['Express 5.2', 'Web Framework'],
    ['TypeScript 5.6', 'Type-safe JavaScript'],
    ['pg 8.22', 'PostgreSQL Client Driver'],
    ['bcrypt 6.0', 'Password Hashing'],
    ['dotenv 17.4', 'Environment Variable Management'],
    ['cors 2.8', 'CORS Middleware'],
    ['multer 2.0', 'File Upload Handling'],
    ['nodemailer 9.0', 'Email Sending'],
    ['tsx 4.19', 'TypeScript Execution (Dev)'],
]
make_table(doc, ['Technology', 'Purpose'], be_stack)

doc.add_paragraph()
doc.add_heading('E. Environment Configuration Reference', level=2)

env_config = [
    ['PORT', '5000', 'Backend server listen port'],
    ['DB_USER', 'postgres', 'PostgreSQL authentication username'],
    ['DB_PASSWORD', '(required)', 'PostgreSQL authentication password'],
    ['DB_HOST', 'db', 'Database hostname (use "db" within Docker network)'],
    ['DB_PORT', '5432', 'PostgreSQL port'],
    ['DB_NAME', 'fieldsync_db', 'Database name'],
    ['JWT_SECRET', '(required)', 'Secret key for JWT token signing and verification'],
    ['EMAIL_USER', '(required)', 'SMTP email address for outgoing notifications'],
    ['EMAIL_PASS', '(required)', 'SMTP email password or application-specific password'],
]
make_table(doc, ['Variable', 'Default', 'Description'], env_config)

# =============================================
# SAVE
# =============================================
output_path = 'C:\\Users\\nebi\\Desktop\\mongoreact\\fieldsync\\FieldSync_System_Design_Document_v2.docx'
doc.save(output_path)
print(f'System Design Document saved to: {output_path}')
