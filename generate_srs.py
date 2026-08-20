from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(30, 58, 95)
    hs.font.name = 'Calibri'

NAVY = '1E3A5F'

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells_data, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        if header:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_shading(cell, NAVY)
    return row

def make_table(doc, headers, rows):
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, headers, header=True)
    for row_data in rows:
        add_table_row(table, row_data)
    return table

# =============================================
# COVER PAGE
# =============================================
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FieldSync')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(30, 58, 95)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Software Requirements Specification (SRS)')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(79, 195, 247)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('A Comprehensive Description of Functionality,\nRequirements, and Constraints')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Version 1.0.0\n{datetime.date.today().strftime("%B %Y")}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 116, 139)

doc.add_page_break()

# =============================================
# TABLE OF CONTENTS
# =============================================
doc.add_heading('Table of Contents', level=1)
for item in [
    '1. Introduction',
    '2. Overall Description',
    '3. Functional Requirements',
    '4. Non-Functional Requirements',
    '5. External Interface Requirements',
    '6. Assumptions and Dependencies',
    '7. Appendices',
]:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(4)

doc.add_page_break()

# =============================================
# 1. INTRODUCTION
# =============================================
doc.add_heading('1. Introduction', level=1)

make_table(doc, ['Field', 'Details'], [
    ['Project Name', 'FieldSync'],
    ['Prepared by', 'FieldSync Development Team'],
    ['Version', '1.0.0'],
    ['Date', datetime.date.today().strftime('%Y-%m-%d')],
    ['Document Purpose', 'This document provides a comprehensive description of the intended purpose, '
     'functionality, and requirements for the FieldSync software system. It serves as a communication '
     'tool between stakeholders, developers, and testers, ensuring a shared understanding of both '
     'functional and non-functional expectations.'],
    ['Intended Audience', 'Project stakeholders, software developers, QA testers, system administrators, '
     'and end-user representatives (Field Officers, Supervisors, Managers).'],
])

# =============================================
# 2. OVERALL DESCRIPTION
# =============================================
doc.add_paragraph()
doc.add_heading('2. Overall Description', level=1)

doc.add_heading('2.1 Product Perspective', level=2)
doc.add_paragraph(
    'FieldSync is a self-contained, cloud-ready field workforce management and citizen registration '
    'platform built as a three-tier web application. It operates as an independent system that '
    'does not rely on external software products for its core functionality. The system is designed '
    'to serve government organizations that deploy field officers for citizen registration programs '
    '(e.g., National ID enrollment) and require robust offline capabilities for areas with '
    'limited or no internet connectivity.'
)
doc.add_paragraph(
    'The product consists of three primary layers: a React-based Progressive Web Application (PWA) '
    'frontend that provides a responsive, installable interface; a Node.js/Express REST API backend '
    'written in TypeScript following MVC architecture; and a PostgreSQL 16 relational database for '
    'persistent data storage. All components are containerized using Docker and orchestrated with '
    'Docker Compose for simplified deployment and scaling.'
)
doc.add_paragraph(
    'A distinguishing feature of FieldSync is its offline-first architecture. Using Dexie.js '
    '(an IndexedDB wrapper), the frontend stores citizen registrations, reports, and pending '
    'operations locally when the device is offline. A built-in synchronization service automatically '
    'queues these operations and replays them against the backend API when connectivity is restored, '
    'ensuring zero data loss in low-connectivity environments.'
)

doc.add_heading('2.2 Product Functions', level=2)
functions = [
    'User Authentication and Authorization: Secure login with email/password (bcrypt + JWT), '
    'role-based access control for three user roles (Manager, Supervisor, Field Officer), '
    'and session management.',
    'Citizen Registration: Comprehensive registration forms for National ID enrollment capturing '
    'biographic data, contact information, address hierarchy (region/district/village), occupation, '
    'marital status, document type (National ID, Birth Certificate, Passport), biometric collection '
    'status, and GPS coordinates.',
    'Field Reporting: Creation, submission, and management of field activity reports with '
    'photo/document attachments, GPS tagging, and offline persistence with automatic synchronization.',
    'Task Management: Supervisor-initiated task assignment to field officers with priority levels, '
    'due dates, completion tracking, and notification alerts.',
    'Team Management: Supervisor tools for monitoring team composition, officer activities, '
    'and performance metrics.',
    'Permission Workflow: Request and approval system for travel and equipment with '
    'multi-level approval routing.',
    'Supervisor Reports: Automated generation of team performance summaries and operational '
    'metrics for supervisory review.',
    'Screen Time Tracking: Application usage monitoring for field devices, recording session '
    'durations and device information.',
    'Identity Verification: Periodic identity verification prompts for field officers to '
    'confirm active duty status with compliance tracking.',
    'Alerts and Messaging: Internal notification and messaging system for team communications, '
    'task notifications, and system alerts.',
    'Analytics Dashboard: Manager-level dashboards with trend charts, registration statistics, '
    'geographic distribution maps, and performance analytics.',
    'Audit Logging: Comprehensive, append-only audit trail of all significant system actions '
    'for accountability and regulatory compliance.',
    'Offline Synchronization: Client-side data queuing in IndexedDB with automatic replay and '
    'conflict resolution upon reconnection.',
    'Multi-Language Support: Full interface localization in English, Amharic, Tigrinya, and Oromo '
    'with persistent language preference.',
    'Progressive Web App: Installable PWA with service worker caching for fast access and '
    'enhanced offline functionality.',
    'Theme Customization: Dark and light theme toggle for user comfort across different '
    'lighting conditions.',
]
for f in functions:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('2.3 User Characteristics', level=2)
doc.add_paragraph(
    'FieldSync is designed for three distinct user roles with varying levels of technical '
    'proficiency and system access:'
)
make_table(doc, ['Role', 'Description', 'Technical Proficiency', 'Primary Tasks'], [
    ['Field Officer',
     'Frontline government workers deployed in the field to register citizens and submit activity reports.',
     'Basic -- comfortable using smartphones and web browsers, may have limited computer literacy.',
     'Register citizens, submit field reports, request permissions, view tasks, respond to verification prompts, communicate via alerts.'],
    ['Supervisor',
     'Team leads who oversee a group of field officers, assign tasks, review reports, and manage team operations.',
     'Moderate -- familiar with web applications and basic data analysis.',
     'Assign and track tasks, manage team members, review and approve permissions, generate supervisor reports, monitor screen time, verify officer activities.'],
    ['Manager',
     'Senior administrators with full system access who make strategic decisions based on data analytics.',
     'Advanced -- proficient with web-based management tools and data analysis.',
     'Manage all users and roles, view analytics dashboards, access all reports and citizen data, review audit logs, configure system settings.'],
])

doc.add_heading('2.4 Constraints', level=2)
constraints = [
    'The system requires Docker Engine 20.10+ and Docker Compose v2+ for deployment; no native '
    'bare-metal installation is supported.',
    'Offline data storage is limited by the browser\'s IndexedDB quota, which varies by browser '
    'and operating system (typically 50 MB to several GB).',
    'All API communication uses HTTP/HTTPS REST; no WebSocket or real-time push notification '
    'infrastructure is included in the current version.',
    'The frontend is served via Nginx in production; server-side rendering is not supported.',
    'File uploads (photos, documents) are limited by server disk space, backed by Docker volumes.',
    'Multi-language support covers four languages (English, Amharic, Tigrinya, Oromo); adding '
    'new languages requires manual translation of all UI strings.',
    'The system does not support multi-tenancy; each deployment serves a single organizational unit.',
    'Biometric data capture is tracked (collection status) but actual biometric processing is '
    'handled by external systems not included in this product.',
    'GPS features require the client device to have location services enabled and browser '
    'location permissions granted.',
    'Email notifications require a configured SMTP server; notification delivery depends on '
    'the availability of the email service.',
    'The PostgreSQL database is deployed within Docker; for high-availability production '
    'deployments, an external managed database service is recommended.',
    'The system targets a minimum of 100 concurrent users on the default Docker Compose '
    'deployment; higher loads require infrastructure scaling.',
]
for c in constraints:
    doc.add_paragraph(c, style='List Bullet')

# =============================================
# 3. FUNCTIONAL REQUIREMENTS
# =============================================
doc.add_heading('3. Functional Requirements', level=1)

functional_reqs = [
    ['FR-AUTH-001', 'User Login: The system shall allow users to log in using email and password credentials.', 'High'],
    ['FR-AUTH-002', 'Password Hashing: The system shall hash all passwords using bcrypt with a minimum of 10 salt rounds before storage.', 'High'],
    ['FR-AUTH-003', 'JWT Token Issuance: The system shall issue a JWT token upon successful authentication for use in subsequent API requests.', 'High'],
    ['FR-AUTH-004', 'Password Change: The system shall allow authenticated users to change their password after verifying the current password.', 'Medium'],
    ['FR-AUTH-005', 'Account Status Check: The system shall reject login attempts from accounts with inactive status.', 'High'],
    ['FR-AUTH-006', 'Session Management: The system shall maintain user sessions via JWT tokens with appropriate expiration.', 'High'],

    ['FR-ROLE-001', 'Role-Based Access: The system shall enforce three roles (Manager, Supervisor, Field Officer) with distinct permissions.', 'High'],
    ['FR-ROLE-002', 'Navigation Adaptation: The frontend shall display sidebar navigation items based on the authenticated user\'s role.', 'High'],
    ['FR-ROLE-003', 'API Authorization: The backend shall restrict API endpoint access based on the requesting user\'s role.', 'High'],

    ['FR-CIT-001', 'Citizen Registration: Field Officers shall be able to register citizens with first name, last name, date of birth, gender, phone number, email, region, district, village, address, occupation, marital status, document type, document number, and biometric status.', 'High'],
    ['FR-CIT-002', 'GPS Capture: The system shall automatically capture GPS coordinates during citizen registration.', 'High'],
    ['FR-CIT-003', 'Citizen Search: The system shall allow users to search and filter registered citizens by name, region, and document type.', 'Medium'],
    ['FR-CIT-004', 'Offline Registration: Field Officers shall be able to register citizens while offline, with data stored locally and synced upon reconnection.', 'High'],

    ['FR-RPT-001', 'Report Creation: Field Officers shall be able to create field reports with title, description, and optional photo/document attachments.', 'High'],
    ['FR-RPT-002', 'Report Submission: Reports shall be submitted to the backend and stored with officer reference, timestamp, and sync status.', 'High'],
    ['FR-RPT-003', 'Offline Report Storage: Reports created while offline shall be stored in IndexedDB and synced automatically when online.', 'High'],
    ['FR-RPT-004', 'Report Viewing: Supervisors and Managers shall be able to view all reports; Field Officers shall see only their own reports.', 'Medium'],

    ['FR-TASK-001', 'Task Creation: Supervisors shall be able to create tasks with title, description, priority level, and due date.', 'High'],
    ['FR-TASK-002', 'Task Assignment: Supervisors shall be able to assign tasks to one or more Field Officers.', 'High'],
    ['FR-TASK-003', 'Task Status Tracking: The system shall track task status (pending, in-progress, completed) and allow officers to update completion.', 'High'],
    ['FR-TASK-004', 'Task Alerts: The system shall send an alert notification to assigned officers when a new task is created.', 'Medium'],

    ['FR-PERM-001', 'Permission Requests: Field Officers shall be able to submit permission requests (travel, equipment) with reason and date range.', 'Medium'],
    ['FR-PERM-002', 'Permission Approval: Supervisors shall be able to approve or reject permission requests with status tracking.', 'Medium'],
    ['FR-PERM-003', 'Permission Notifications: The system shall notify the requesting officer when their permission request is approved or rejected.', 'Low'],

    ['FR-TEAM-001', 'Team Overview: Supervisors shall be able to view a list of their assigned team members with profile details and status.', 'Medium'],
    ['FR-TEAM-002', 'Team Performance: Supervisors shall be able to view team activity summaries and performance metrics.', 'Medium'],

    ['FR-RPT-SUP-001', 'Supervisor Report Generation: Supervisors shall be able to generate reports summarizing team activities and performance for a specified period.', 'Medium'],
    ['FR-RPT-SUP-002', 'Report Viewing: Managers shall be able to view all supervisor reports across teams.', 'Medium'],

    ['FR-SCREN-001', 'Screen Time Logging: The system shall record application session duration for each user.', 'Low'],
    ['FR-SCREN-002', 'Screen Time Reports: Supervisors shall be able to view screen time data for their team members.', 'Low'],

    ['FR-VERIF-001', 'Verification Prompts: The system shall periodically prompt users to confirm their identity.', 'Medium'],
    ['FR-VERIF-002', 'Verification Compliance: The system shall track verification response times and notify supervisors of missed verifications.', 'Medium'],

    ['FR-ALERT-001', 'Alert Creation: The system shall support internal alert and message creation between users.', 'Medium'],
    ['FR-ALERT-002', 'Alert Notifications: The system shall display unread alert count as a badge in the navigation.', 'Medium'],
    ['FR-ALERT-003', 'Alert Read Status: The system shall track whether alerts have been read by recipients.', 'Low'],

    ['FR-USR-001', 'User Management: Managers shall be able to create, edit, and deactivate user accounts.', 'High'],
    ['FR-USR-002', 'Role Assignment: Managers shall be able to assign roles (Manager, Supervisor, Field Officer) to users.', 'High'],
    ['FR-USR-003', 'Profile Management: Users shall be able to view and edit their profile, including uploading a profile photo.', 'Medium'],
    ['FR-USR-004', 'User Listing: Managers shall be able to view all users; Supervisors shall view their team members.', 'Medium'],

    ['FR-LOC-001', 'Location Hierarchy: The system shall provide cascading location selection (region > district > village) for registration and reporting.', 'Medium'],
    ['FR-LOC-002', 'Location Data: The system shall store and serve geographic reference data from the locations table.', 'Low'],

    ['FR-AUDIT-001', 'Audit Logging: The system shall log all significant actions (create, update, delete) with user, action type, entity, and timestamp.', 'High'],
    ['FR-AUDIT-002', 'Audit Viewing: Managers shall be able to view and filter audit logs.', 'Medium'],

    ['FR-SYNC-001', 'Offline Queue: The system shall queue all write operations in IndexedDB when the device is offline.', 'High'],
    ['FR-SYNC-002', 'Auto-Sync: The system shall automatically replay queued operations against the API when connectivity is restored.', 'High'],
    ['FR-SYNC-003', 'Sync Status: The system shall display real-time sync status and pending item count in the UI.', 'High'],
    ['FR-SYNC-004', 'Conflict Resolution: The system shall handle data conflicts during sync with timestamp-based resolution.', 'Medium'],

    ['FR-I18N-001', 'Language Selection: The system shall allow users to switch between English, Amharic, Tigrinya, and Oromo.', 'Medium'],
    ['FR-I18N-002', 'Persistent Language: The system shall remember and restore the user\'s language preference across sessions.', 'Low'],

    ['FR-PWA-001', 'PWA Installation: The system shall be installable as a Progressive Web App on supported devices.', 'Low'],
    ['FR-PWA-002', 'Service Worker Caching: The system shall cache static assets via service worker for offline access.', 'Medium'],

    ['FR-THEME-001', 'Theme Toggle: The system shall provide dark and light theme options with a toggle switch.', 'Low'],
]
make_table(doc, ['ID', 'Description', 'Priority'], functional_reqs)

# =============================================
# 4. NON-FUNCTIONAL REQUIREMENTS
# =============================================
doc.add_paragraph()
doc.add_heading('4. Non-Functional Requirements', level=1)

nfr = [
    ['Performance', 'All API endpoints shall respond within 500ms at the 95th percentile under normal load.'],
    ['Performance', 'The initial page load shall complete within 3 seconds on a 10 Mbps broadband connection.'],
    ['Performance', 'The system shall support a minimum of 100 concurrent users without response time degradation.'],
    ['Performance', 'The Sync Service shall sync 50+ pending records within 60 seconds upon reconnection.'],
    ['Performance', 'Database CRUD queries shall execute within 100ms; complex analytics queries within 2 seconds.'],
    ['Reliability', 'The system shall maintain 99.5% uptime during operational hours (8 AM - 5 PM).'],
    ['Reliability', 'The system shall transition to offline mode within 2 seconds of network disconnection with zero data loss.'],
    ['Reliability', 'All data operations shall be atomic; partial writes shall not corrupt the database.'],
    ['Reliability', 'Docker health checks shall automatically restart failed containers within 30 seconds.'],
    ['Usability', 'The interface shall support screen resolutions from 1024x768 (desktop) to mobile viewports.'],
    ['Usability', 'The system shall provide multilingual support in English, Amharic, Tigrinya, and Oromo.'],
    ['Usability', 'Form validation shall provide clear, inline error messages for all required fields.'],
    ['Usability', 'The system shall provide a guided login experience with visible online/offline status indicators.'],
    ['Security', 'All passwords shall be hashed using bcrypt with a minimum of 10 salt rounds.'],
    ['Security', 'All API endpoints (except login) shall require JWT Bearer token authentication.'],
    ['Security', 'CORS shall be configured to restrict API access to authorized origins only.'],
    ['Security', 'Sensitive configuration (database credentials, JWT secret) shall be stored in environment variables, not in source code.'],
    ['Security', 'All significant system actions shall be logged in an append-only audit trail.'],
    ['Scalability', 'The Docker Compose deployment shall be horizontally scalable by increasing container replicas.'],
    ['Scalability', 'The PostgreSQL database shall support connection pooling for efficient resource utilization.'],
    ['Scalability', 'File storage shall use Docker volumes for persistence and portability.'],
    ['Compatibility', 'The frontend shall be compatible with Chrome 90+, Firefox 88+, Edge 90+, and Safari 14+.'],
    ['Compatibility', 'The PWA shall install and function on Android 8+, iOS 14+, Windows 10+, and macOS 12+.'],
    ['Maintainability', 'The backend shall follow MVC architecture with clear separation of routes, controllers, models, and utils.'],
    ['Maintainability', 'All backend code shall be written in TypeScript with strict type checking enabled.'],
    ['Maintainability', 'The frontend shall use consistent component structure with shared hooks, contexts, and services.'],
    ['Data Integrity', 'Database migrations shall be version-controlled and applied automatically on container startup.'],
    ['Data Integrity', 'Form inputs shall be validated both client-side (Zod schemas) and server-side before persistence.'],
    ['Data Integrity', 'File uploads shall be validated for file type and size before server-side storage.'],
    ['Portability', 'The entire system shall deploy via a single "docker compose up -d --build" command.'],
    ['Portability', 'No platform-specific dependencies shall exist outside the Docker container environment.'],
]
make_table(doc, ['Category', 'Requirement'], nfr)

# =============================================
# 5. EXTERNAL INTERFACE REQUIREMENTS
# =============================================
doc.add_paragraph()
doc.add_heading('5. External Interface Requirements', level=1)

ext_ifs = [
    ['User Interface (Web Browser)',
     'The system shall provide a responsive web interface accessible via modern browsers. The UI shall include '
     'a login page, role-based dashboard, sidebar navigation, forms with validation, data tables, charts, '
     'notification indicators, and theme toggle. The interface shall adapt to desktop and mobile screen sizes.'],
    ['REST API Interface',
     'The system shall expose a RESTful HTTP API using JSON for all data operations. The API shall follow '
     'RESTful conventions (GET, POST, PUT, DELETE) with consistent response formats. All endpoints except '
     '/api/auth/login shall require JWT authentication via the Authorization header.'],
    ['PostgreSQL Database Interface',
     'The backend shall connect to PostgreSQL 16 via the node-postgres (pg) driver with connection pooling. '
     'The connection shall use TCP on port 5432 within the Docker network. All queries shall use parameterized '
     'statements to prevent SQL injection.'],
    ['SMTP Email Interface',
     'The system shall send email notifications using Nodemailer via SMTP/TLS. Configuration shall be provided '
     'through EMAIL_USER and EMAIL_PASS environment variables. Email shall be sent for task notifications, '
     'permission status changes, and other system-triggered events.'],
     ['Browser Geolocation API',
      'The frontend shall access the browser\'s Geolocation API to capture GPS coordinates during citizen '
      'registration. The application shall handle permission denial gracefully with '
      'appropriate user feedback.'],
    ['Browser IndexedDB / Dexie.js',
     'The frontend shall use Dexie.js as a wrapper around IndexedDB for offline data storage. The storage '
     'layer shall support CRUD operations on citizens, reports, and pending sync items with automatic '
     'synchronization upon reconnection.'],
    ['Browser Camera / MediaDevices API',
     'The frontend shall optionally access the device camera for profile photo capture and citizen document '
     'scanning via the MediaDevices API. Camera access shall require explicit user permission.'],
    ['Docker Engine Interface',
     'The system shall be deployed and managed via Docker Compose. The deployment stack shall define three '
     'services (db, backend, frontend) with health checks, dependency ordering, volume mounts, port mappings, '
     'and environment variable injection.'],
    ['Progressive Web App (Service Worker)',
     'The frontend shall register a service worker via vite-plugin-pwa for asset caching and offline support. '
     'The service worker shall cache static assets (HTML, CSS, JS, images) and provide a fallback page '
     'when the network is unavailable.'],
    ['File System (Uploads)',
     'The backend shall store uploaded files (photos, documents) to a local file system path mapped to a Docker '
     'volume (uploads). Multer shall handle multipart/form-data parsing with configured size limits and '
     'destination paths.'],
]
make_table(doc, ['Interface', 'Description'], ext_ifs)

# =============================================
# 6. ASSUMPTIONS AND DEPENDENCIES
# =============================================
doc.add_paragraph()
doc.add_heading('6. Assumptions and Dependencies', level=1)

assumptions = [
    'Docker Engine (20.10+) and Docker Compose (v2+) are installed and operational on the deployment server.',
    'The deployment environment provides reliable power supply and at least intermittent internet connectivity for initial setup and periodic synchronization.',
    'Field officers have access to modern web browsers on smartphones, tablets, or laptops with GPS and camera capabilities.',
    'An SMTP email server is available and properly configured for outgoing notification emails.',
    'Administrators have sufficient command-line proficiency to manage Docker-based deployments, including viewing logs and restarting services.',
    'The PostgreSQL database running within Docker provides adequate performance for the expected data volumes (10,000+ citizen records, 5,000+ reports).',
    'Browser IndexedDB storage quotas on field devices are sufficient for offline data caching during typical field deployment periods (1-5 days between sync opportunities).',
    'Network infrastructure in deployment regions supports HTTP/HTTPS traffic on standard ports (80, 443, 5001, 30001).',
    'The organizational entity deploying FieldSync has the legal authority to collect and process citizen personal data as required for National ID registration.',
    'End users (Field Officers) will receive basic training on system usage, including login, citizen registration, report creation, and offline mode operation.',
    'The system will be backed up regularly using Docker volume backup procedures or PostgreSQL dump utilities.',
    'Time synchronization (NTP) is maintained across all deployment servers to ensure accurate timestamps in audit logs and data synchronization.',
    'Translation and localization for the four supported languages (English, Amharic, Tigrinya, Oromo) will be provided and reviewed by native speakers.',
    'Profile photos and citizen document images will be of reasonable resolution and file size (under 5 MB per file) to avoid storage exhaustion.',
]
for a in assumptions:
    doc.add_paragraph(a, style='List Bullet')

# =============================================
# 7. APPENDICES
# =============================================
doc.add_heading('7. Appendices', level=1)

doc.add_heading('Appendix A: Glossary', level=2)

glossary = [
    ['API', 'Application Programming Interface -- a set of protocols for building and interacting with software applications.'],
    ['bcrypt', 'A password hashing algorithm designed to be computationally expensive to resist brute-force attacks.'],
    ['CRUD', 'Create, Read, Update, Delete -- the four fundamental database operations.'],
    ['CORS', 'Cross-Origin Resource Sharing -- a security mechanism controlling cross-domain HTTP requests.'],
    ['Dexie.js', 'A JavaScript library providing a simplified wrapper around the browser IndexedDB API.'],
    ['Docker Compose', 'A tool for defining and running multi-container Docker applications.'],
    ['IndexedDB', 'A browser-based NoSQL database for client-side storage of structured data.'],
    ['i18next', 'An internationalization framework providing translation and language detection.'],
    ['JWT', 'JSON Web Token -- a compact token format for securely transmitting information between parties.'],
    ['MVC', 'Model-View-Controller -- an architectural pattern separating data, presentation, and logic.'],
    ['Multer', 'A Node.js middleware for handling file uploads via multipart/form-data.'],
    ['Nginx', 'A high-performance web server used as a reverse proxy and static file server.'],
    ['Nodemailer', 'A Node.js library for sending emails via SMTP.'],
    ['PWA', 'Progressive Web App -- a web application providing native app-like capabilities.'],
    ['RBAC', 'Role-Based Access Control -- access regulation based on assigned user roles.'],
    ['REST', 'Representational State Transfer -- an architectural style for networked applications.'],
    ['SPA', 'Single Page Application -- a web app that loads once and dynamically updates content.'],
    ['Sync Service', 'A client-side module managing offline data queuing and server synchronization.'],
    ['Vite', 'A modern frontend build tool with fast hot module replacement.'],
    ['Zod', 'A TypeScript-first schema validation library for form and API input validation.'],
]
make_table(doc, ['Term', 'Definition'], glossary)

doc.add_paragraph()
doc.add_heading('Appendix B: API Endpoint Summary', level=2)

api_endpoints = [
    ['POST', '/api/auth/login', 'User authentication, returns JWT token'],
    ['POST', '/api/auth/change-password', 'Change user password'],
    ['GET', '/api/users', 'List all users (filtered by role)'],
    ['POST', '/api/users', 'Create a new user account'],
    ['PUT', '/api/users/:id', 'Update user profile'],
    ['DELETE', '/api/users/:id', 'Deactivate a user account'],
    ['GET', '/api/citizens', 'List registered citizens'],
    ['POST', '/api/citizens', 'Register a new citizen'],
    ['GET', '/api/reports', 'List field reports'],
    ['POST', '/api/reports', 'Submit a new field report'],
    ['GET', '/api/tasks', 'List assigned tasks'],
    ['POST', '/api/tasks', 'Create a new task'],
    ['PUT', '/api/tasks/:id', 'Update task status'],
    ['GET', '/api/permissions', 'List permission requests'],
    ['POST', '/api/permissions', 'Submit a permission request'],
    ['PUT', '/api/permissions/:id', 'Approve/reject permission'],
    ['GET', '/api/alerts', 'List alerts and messages'],
    ['POST', '/api/alerts', 'Send a new alert'],
    ['GET', '/api/audit', 'View audit log entries'],
    ['GET', '/api/screentime', 'View screen time data'],
    ['GET', '/api/verification', 'View verification records'],
    ['POST', '/api/verification/respond', 'Respond to verification prompt'],
    ['GET', '/api/locations', 'Get location hierarchy data'],
    ['GET', '/api/supervisor-reports', 'View supervisor reports'],
    ['POST', '/api/supervisor-reports', 'Generate supervisor report'],
    ['POST', '/api/sync', 'Sync offline data to server'],
]
make_table(doc, ['Method', 'Endpoint', 'Description'], api_endpoints)

doc.add_paragraph()
doc.add_heading('Appendix C: Database Tables', level=2)

db_tables = [
    ['users', 'User accounts with credentials, roles, profile data, and status'],
    ['citizens', 'Registered citizen records with biographic data and GPS coordinates'],
    ['reports', 'Field activity reports with attachments and sync status'],
    ['tasks', 'Assigned tasks with priority, due dates, and completion status'],
    ['permissions', 'Permission requests and approval workflow records'],
    ['alerts', 'Internal messaging and notification records'],
    ['audit', 'System action audit trail with user, action, entity, and timestamp'],
    ['screen_time', 'Application session duration and device tracking'],
    ['verification', 'Identity verification prompts and compliance records'],
    ['locations', 'Geographic reference data (regions, districts, villages)'],
    ['supervisor_reports', 'Generated supervisor performance reports'],
]
make_table(doc, ['Table', 'Description'], db_tables)

doc.add_paragraph()
doc.add_heading('Appendix D: Revision History', level=2)
make_table(doc, ['Version', 'Date', 'Changes'], [
    ['1.0.0', datetime.date.today().strftime('%Y-%m-%d'), 'Initial release of the FieldSync SRS document.'],
])

# =============================================
# SAVE
# =============================================
output_path = 'C:\\Users\\nebi\\Desktop\\mongoreact\\fieldsync\\FieldSync_SRS_v2.docx'
doc.save(output_path)
print(f'SRS Document saved to: {output_path}')
